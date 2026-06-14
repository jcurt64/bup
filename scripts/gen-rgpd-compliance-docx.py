#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère 3 documents Word de conformité RGPD pour BUUPP, ANCRÉS sur le système
réel (politique /rgpd publiée + base de données + code) et alignés sur les
référentiels CNIL :

  1. Registre-des-traitements-BUUPP.docx        (RoPA — art. 30 RGPD)
  2. Referentiel-durees-conservation-BUUPP.docx (cycle de vie des données)
  3. AIPD-BUUPP-mise-en-relation-scoring.docx   (AIPD/DPIA — art. 35 RGPD)

Usage : python3 scripts/gen-rgpd-compliance-docx.py
Sortie : docs/rgpd/*.docx

⚠️ Documents de travail : ils décrivent l'existant technique. Les éléments
juridiques (désignation formelle du DPO, validation des bases légales,
arbitrages de durées) doivent être revus/validés par le DPO et le conseil
juridique. Les champs à arbitrer sont marqués « [À COMPLÉTER] ».
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
INK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x47, 0x55, 0x65)
RED = RGBColor(0xB4, 0x53, 0x09)

DATE = "14/06/2026"
EDITEUR = ("Majelink (éditeur de la plateforme BUUPP) — RCS Pau 892 514 167 — "
           "12 Impasse des Étriers, 64140 Lons")
DPO = "Chargé à la protection des données : dp.buupp@buupp.com"


def base_doc():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10)
    n.font.color.rgb = INK
    return doc


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = INDIGO if level <= 1 else INK
    return p


def P(doc, text, italic=False, bold=False, color=None, size=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic, r.bold = italic, bold
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p


def BUL(doc, text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        p.add_run(prefix).bold = True
    p.add_run(text)
    return p


def disclaimer(doc):
    p = doc.add_paragraph()
    r = p.add_run(
        "Document de travail ancré sur le système en production (politique de "
        "confidentialité publiée, base de données, code source). Les arbitrages "
        "juridiques (bases légales, durées, désignation du DPO) doivent être "
        "validés par le DPO / le conseil juridique. Les champs marqués "
        "« [À COMPLÉTER] » requièrent une décision du responsable de traitement.")
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(10)


def cover(doc, title, subtitle):
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BUUPP"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = INDIGO
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(title); r.font.size = Pt(16); r.font.color.rgb = INK
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run(subtitle); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
    doc.add_paragraph()
    m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("Responsable de traitement : ").bold = True
    m.add_run(EDITEUR + "\n")
    m.add_run(DPO + "\n")
    m.add_run("Version : " + DATE)
    doc.add_paragraph()
    disclaimer(doc)


def table(doc, headers, rows, widths=None, font=8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, label in zip(t.rows[0].cells, headers):
        run = c.paragraphs[0].add_run(label); run.bold = True; run.font.size = Pt(font)
    for row in rows:
        cells = t.add_row().cells
        for c, val in zip(cells, row):
            run = c.paragraphs[0].add_run(val); run.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for c, w in zip(row.cells, widths):
                c.width = Inches(w)
    return t


# Sous-traitants (tiré de la politique /rgpd publiée + système vérifié)
SOUS_TRAITANTS = [
    ["Supabase Inc.", "Base de données PostgreSQL + Storage", "UE — Irlande (eu-west-1)*", "Art. 28 — DPA"],
    ["Vercel Inc.", "Hébergement de l'application", "États-Unis (servi en UE)", "SCC (clauses types)"],
    ["Clerk Inc.", "Authentification / identités", "États-Unis", "SCC (clauses types)"],
    ["Stripe Payments Europe Ltd", "Paiements, KYC, virements (Connect)", "UE — Irlande", "RT distinct (PCI-DSS)"],
    ["Brevo (Sendinblue SAS)", "E-mails et SMS transactionnels", "UE — France", "Art. 28 — DPA"],
    ["Cloudflare", "Routage des alias e-mail watermarqués", "Réseau mondial (UE)", "SCC (clauses types)"],
    ["DGFiP", "Transmission fiscale annuelle (DAC7)", "France", "Destinataire légal"],
]


# ════════════════════════════════════════════════════════════════════
# DOC 1 — REGISTRE DES TRAITEMENTS
# ════════════════════════════════════════════════════════════════════
def build_registre():
    doc = base_doc()
    cover(doc, "Registre des activités de traitement",
          "RoPA — article 30 du RGPD")
    doc.add_page_break()

    H(doc, "Informations générales", 1)
    BUL(doc, EDITEUR, "Responsable de traitement : ")
    BUL(doc, "dp.buupp@buupp.com", "DPO / Chargé à la protection des données : ")
    BUL(doc, "Particuliers (« prospects »), professionnels clients, visiteurs (liste d'attente).",
        "Catégories de personnes concernées : ")
    BUL(doc, "Aucun transfert hors UE pour les données de production hébergées chez Supabase (UE). "
             "Certains sous-traitants (Vercel, Clerk, Cloudflare) sont établis aux États-Unis : "
             "encadrement par clauses contractuelles types (SCC).",
        "Transferts hors UE : ")
    P(doc, "* La politique publiée mentionne « Francfort » ; le projet de production est "
           "effectivement en région eu-west-1 (Irlande). À harmoniser entre la politique et ce registre. [À COMPLÉTER]",
      italic=True, size=8.5, color=RED)

    H(doc, "Sous-traitants et destinataires (art. 28 / 30.1.d)", 1)
    table(doc, ["Sous-traitant", "Rôle", "Localisation", "Encadrement"],
          SOUS_TRAITANTS, font=8.5)
    P(doc, "Les professionnels destinataires d'une sollicitation acceptée deviennent "
           "responsables de traitement distincts pour les seules coordonnées révélées. "
           "Tenir à jour les DPA (art. 28) avec chaque sous-traitant. [À COMPLÉTER : dates de signature des DPA]",
      italic=True, size=9)

    H(doc, "Fiches de traitement", 1)
    P(doc, "Mesures de sécurité communes à tous les traitements : hébergement UE, "
           "chiffrement en transit (TLS) et au repos, authentification Clerk + middleware, "
           "Row Level Security (tables d'audit sans policy = service_role uniquement), "
           "rate-limiting, pseudonymisation à la lecture, journalisation des accès "
           "(fail-closed + append-only), anti-fraude (unicité IBAN/téléphone/rôle, honeypots).",
      size=9)

    # (titre, finalité, base légale, personnes, données, destinataires, durée, particularités)
    fiches = [
        ("T1 — Gestion des comptes utilisateurs",
         "Création / gestion du compte, vérification e-mail et téléphone, sécurité.",
         "Exécution du contrat (art. 6.1.b) ; vérification téléphone : intérêt légitime / sécurité.",
         "Prospects, professionnels.",
         "Authentification (id Clerk, e-mail, méthode de connexion) ; Palier 1 identité prospect "
         "(prénom, nom, e-mail, téléphone vérifié SMS, date de naissance, genre, nationalité) ; "
         "profil pro (raison sociale, SIREN/SIRET, forme juridique, RCS, adresse).",
         "Personnels habilités Majelink ; Clerk ; Brevo (SMS/e-mail) ; Stripe (KYC pro).",
         "Durée du compte ; suppression des données identifiantes à la clôture (cf. référentiel).",
         "Téléphone vérifié par OTP ; unicité téléphone/IBAN/e-mail (anti-doublon)."),

        ("T2 — Données de profil par paliers (2 à 5)",
         "Enrichissement facultatif du profil pour la mise en relation ciblée.",
         "Consentement (art. 6.1.a) — paliers 2 à 5 facultatifs.",
         "Prospects.",
         "Localisation (adresse, CP, ville, région, logement, mobilité) ; style de vie "
         "(famille, véhicule, sport, animaux) ; données professionnelles (poste, revenus, "
         "statut, secteur) ; patrimoine & projets (propriétaire/locataire, épargne, projet "
         "immobilier, succession, création d'entreprise).",
         "Personnels habilités Majelink ; transmis pseudonymisés aux pros (cf. T4).",
         "Durée du compte ; effacement à la demande / au retrait du consentement.",
         "Données financières (revenus, patrimoine) = données « hautement personnelles » → AIPD."),

        ("T3 — Mise en relation rémunérée (double consentement)",
         "Ciblage par palier, acceptation explicite des sollicitations, calcul de la "
         "rémunération, traçabilité du contrat.",
         "Exécution du contrat (art. 6.1.b) ; acceptation d'une sollicitation : consentement (6.1.a).",
         "Prospects, professionnels.",
         "Données de relation : sollicitations, acceptations/refus, motif du pro, date de "
         "décision, code d'authentification de campagne, évaluation post-contact.",
         "Prospect, professionnel concerné, personnels Majelink.",
         "Durée du compte / archivage comptable des transactions (10 ans).",
         "Aucune donnée transmise sans accord explicite ; séquestre jusqu'à clôture."),

        ("T4 — Pseudonymisation et transmission aux professionnels",
         "Transmettre au pro un profil exploitable mais non directement identifiant.",
         "Exécution du contrat (art. 6.1.b) ; consentement à la sollicitation (6.1.a).",
         "Prospects.",
         "Profil pseudonymisé : nom masqué, e-mail → alias watermarqué, date de naissance → "
         "tranche d'âge, code postal → département, adresse/poste/revenus → supprimés.",
         "Professionnel destinataire (RT distinct pour les coordonnées révélées).",
         "Cf. T3 ; journal des révélations conservé 24 mois.",
         "Réversible par BUUPP seul ; chaque révélation journalisée (fail-closed + append-only)."),

        ("T5 — Calcul du BUUPP Score et segmentation",
         "Noter la qualité de profil et permettre la segmentation d'audience par les pros.",
         "Intérêt légitime (art. 6.1.f) — valorisation du profil et pertinence du ciblage. [À COMPLÉTER : balance des intérêts]",
         "Prospects.",
         "Score (0-1000) dérivé des paliers ; facettes de segmentation (score, région, distance, "
         "logement, statut pro, foyer, véhicule, animaux).",
         "Professionnels (vue agrégée + segments) ; personnels Majelink.",
         "Durée du compte.",
         "Profilage/scoring + croisement de données → relève de l'AIPD (art. 35)."),

        ("T6 — Gestion financière et rémunération",
         "Recharge des portefeuilles pro, séquestre, débit, paiement et retrait sur IBAN prospect.",
         "Exécution du contrat (art. 6.1.b).",
         "Prospects, professionnels.",
         "Identifiant client Stripe (token), IBAN prospect (retraits), historique des "
         "transactions wallet, factures.",
         "Stripe (paiements/virements), personnels Majelink (compta).",
         "IBAN : durée du compte ; pièces comptables : 10 ans.",
         "Aucune donnée carte complète stockée par BUUPP (déléguée à Stripe, PCI-DSS)."),

        ("T7 — Facturation et obligations fiscales (DAC7)",
         "Émission de factures et déclaration DGFiP au-delà des seuils (2 000 € ou 30 transactions/an).",
         "Obligation légale (art. 6.1.c — CGI art. 242 bis, directive UE DAC7).",
         "Prospects (au-delà du seuil), professionnels.",
         "Cumul annuel des gains, nombre de transactions, attestations DGFiP, factures.",
         "DGFiP ; personnels Majelink ; comptable.",
         "Fiscal DAC7 : 6 ans ; comptabilité : 10 ans.",
         "—"),

        ("T8 — Communication transactionnelle (e-mail / SMS)",
         "Confirmations de sollicitation, alertes d'encaissement, notifications, OTP téléphone.",
         "Exécution du contrat (art. 6.1.b) ; SMS OTP : sécurité.",
         "Prospects, professionnels.",
         "E-mail, téléphone, contenu transactionnel.",
         "Brevo (UE) ; personnels Majelink.",
         "Durée du compte ; e-mails Pro→Prospect : 12 mois.",
         "Domaine authentifié SPF/DKIM/DMARC."),

        ("T9 — Broadcasts pros → prospects + mesure d'ouverture",
         "Diffusion médiée de messages à un segment ; mesure d'ouverture (pixel).",
         "Exécution du contrat (6.1.b) ; pixel de tracking : consentement (6.1.a).",
         "Prospects.",
         "Sujet/corps des messages, identifiant destinataire opaque, statut d'ouverture (pixel).",
         "Brevo ; professionnel émetteur ; personnels Majelink.",
         "Tracking pixel : 13 mois ; e-mails : 12 mois.",
         "Le pixel n'est inséré qu'avec consentement explicite."),

        ("T10 — Prévention de la fraude et journalisation des accès",
         "Anti-doublon, exclusivité de rôle, détection d'abus, traçabilité des révélations.",
         "Intérêt légitime (art. 6.1.f) — sécurité du service et lutte contre la fraude.",
         "Prospects, professionnels.",
         "Contraintes d'unicité (IBAN, téléphone E.164, e-mail, rôle), honeypots, alias "
         "watermarqués, journal des révélations (qui/quoi/quand), clics de contact.",
         "Personnels habilités Majelink (conformité).",
         "Journal des révélations : 24 mois ; clics de contact : 24 mois.",
         "Journal verrouillé (append-only) + écriture fail-closed."),

        ("T11 — Liste d'attente (waitlist) et anti-bot",
         "Inscription en liste d'attente avant lancement, protection anti-bot.",
         "Consentement (6.1.a) pour l'inscription ; intérêt légitime (6.1.f) pour l'anti-bot.",
         "Visiteurs / futurs prospects.",
         "E-mail, hash IP (SHA-256 salé, non réversible), user-agent.",
         "Personnels Majelink.",
         "IP hashées : 12 mois ; e-mails de prospection : cf. référentiel (3 ans CNIL).",
         "Honeypot anti-bot ; IP pseudonymisée par hachage salé."),

        ("T12 — Gestion des demandes RGPD / relation DPO",
         "Réception et traitement des demandes d'exercice de droits.",
         "Obligation légale (art. 6.1.c — art. 12 à 22 RGPD).",
         "Prospects, professionnels, anciens utilisateurs, tiers.",
         "Identité du demandeur, objet de la demande, échanges, justificatif d'identité.",
         "DPO / personnels habilités Majelink.",
         "Le temps du traitement de la demande + délai de preuve (cf. référentiel).",
         "Formulaire DPO public (honeypot + consentement)."),

        ("T13 — Analyse interne et statistiques agrégées",
         "Pilotage du service via des agrégats anonymisés ; back-office d'administration.",
         "Intérêt légitime (art. 6.1.f).",
         "Prospects, professionnels (agrégés).",
         "Agrégats anonymisés ; journal d'événements administrateur (admin_events).",
         "Personnels habilités Majelink.",
         "Logs techniques/sécurité : 12 mois.",
         "Mesure d'audience web via Vercel Analytics — sans cookie."),
    ]

    for (titre, fin, base, pers, donnees, dest, duree, part) in fiches:
        H(doc, titre, 2)
        BUL(doc, fin, "Finalité : ")
        BUL(doc, base, "Base légale : ")
        BUL(doc, pers, "Personnes concernées : ")
        BUL(doc, donnees, "Catégories de données : ")
        BUL(doc, dest, "Destinataires / sous-traitants : ")
        BUL(doc, duree, "Durée de conservation : ")
        if part != "—":
            BUL(doc, part, "Particularités / mesures : ")

    out = "docs/rgpd/Registre-des-traitements-BUUPP.docx"
    doc.save(out)
    return out


# ════════════════════════════════════════════════════════════════════
# DOC 2 — RÉFÉRENTIEL DES DURÉES DE CONSERVATION
# ════════════════════════════════════════════════════════════════════
def build_referentiel():
    doc = base_doc()
    cover(doc, "Référentiel des durées de conservation",
          "Cycle de vie des données — base active / archivage / suppression")
    doc.add_page_break()

    H(doc, "Méthode", 1)
    P(doc, "Conformément au principe de limitation de la conservation (art. 5.1.e RGPD) et "
           "à la méthode CNIL du cycle de vie des données, chaque catégorie est conservée :")
    BUL(doc, "en BASE ACTIVE le temps strictement nécessaire à la finalité ;")
    BUL(doc, "puis, le cas échéant, en ARCHIVAGE INTERMÉDIAIRE (accès restreint) pour répondre "
             "à une obligation légale ou à un besoin de preuve (contentieux) ;")
    BUL(doc, "puis SUPPRESSION définitive ou anonymisation.")
    P(doc, "Les durées ci-dessous reprennent la politique de confidentialité publiée de BUUPP "
           "et les recommandations / obligations applicables (CNIL, Code de commerce, CGI, LCEN). "
           "Les durées non encore arbitrées sont marquées [À COMPLÉTER].", size=9, italic=True)

    H(doc, "Tableau des durées", 1)
    rows = [
        ["Compte & identité (prospect/pro)",
         "Durée du compte (relation contractuelle)",
         "Suppression immédiate des données identifiantes à la clôture",
         "Nécessité au contrat (CNIL : pas au-delà du nécessaire)"],
        ["Paliers 2-5 (profil facultatif)",
         "Durée du compte tant que le consentement est maintenu",
         "Effacement à la demande / au retrait du consentement",
         "Consentement (art. 6.1.a) ; minimisation"],
        ["Transactions / pièces comptables / factures",
         "Durée de la relation",
         "Archivage 10 ans",
         "Art. L.123-22 Code de commerce (compta) ; LPF L.102 B"],
        ["Données fiscales (DAC7)",
         "Exercice en cours",
         "6 ans à compter de la fin de l'exercice",
         "CGI art. 242 bis ; directive UE DAC7"],
        ["IBAN prospect (rémunération)",
         "Durée du compte (besoin pour les retraits)",
         "Supprimé à la clôture ; preuve de paiement via pièce comptable",
         "CNIL : données bancaires conservées le temps de la prestation"],
        ["Identifiant client Stripe (token)",
         "Durée du compte",
         "Supprimé à la clôture",
         "Donnée carte gérée par Stripe (RT distinct, PCI-DSS)"],
        ["Journal d'audit des révélations (pro_contact_reveals)",
         "Accès en base active",
         "24 mois puis purge automatique quotidienne",
         "Accountability (art. 5.2) ; durée validée — à confirmer DPO"],
        ["Actions de contact / click-to-call (audit)",
         "—",
         "24 mois après l'événement",
         "Preuve en cas de litige / signalement"],
        ["Logs techniques et de sécurité",
         "—",
         "12 mois maximum",
         "CNIL recommande ~6 mois ; LCEN (données de connexion) jusqu'à 1 an"],
        ["E-mails Pro → Prospect (contenu)",
         "—",
         "12 mois",
         "Audit anti-spam / détection d'abus (intérêt légitime)"],
        ["Tracking pixel des broadcasts",
         "—",
         "13 mois après envoi",
         "Consentement ; aligné durée traceur CNIL"],
        ["IP hashées (waitlist)",
         "—",
         "12 mois",
         "Anti-bot (intérêt légitime) ; IP pseudonymisée (hash salé)"],
        ["E-mails de prospection (waitlist/contact non client)",
         "Jusqu'au lancement / à la réponse",
         "3 ans à compter du dernier contact",
         "Recommandation CNIL — prospection commerciale"],
        ["Cookies / traceurs non essentiels",
         "Dépôt soumis à consentement",
         "Traceur ≤ 13 mois ; informations collectées ≤ 25 mois",
         "Délibérations CNIL 2020-091/092"],
        ["Mesure d'audience (Vercel Analytics)",
         "—",
         "Sans cookie ; agrégats",
         "Exemptée de consentement si conforme aux conditions CNIL"],
        ["Demandes d'exercice de droits (DPO)",
         "Le temps du traitement de la demande",
         "Durée de preuve (ex. 1 an au-delà du dernier échange)",
         "Obligation légale (art. 12 RGPD) ; [À COMPLÉTER]"],
        ["Journal d'événements administrateur (admin_events)",
         "—",
         "12 mois (logs)",
         "Sécurité / supervision (intérêt légitime)"],
    ]
    table(doc, ["Catégorie de données", "Base active", "Archivage / suppression",
                "Justification & référence"],
          rows, font=8)

    H(doc, "Mises en œuvre techniques de la purge", 1)
    BUL(doc, "Journal des révélations : purge quotidienne automatisée "
             "(lib/pro/reveals-retention.ts, cron /api/admin/digest), durée 24 mois.")
    BUL(doc, "Bascule de durées de conservation CNIL (réinitialisations) : pilotée par cron "
             "+ configuration (lib/cnil/bascule.ts).")
    BUL(doc, "Effacement à la clôture de compte : suppression des données identifiantes, "
             "conservation des transactions en archive comptable (10 ans).")
    P(doc, "Points à finaliser : automatiser la purge des e-mails de prospection à 3 ans, "
           "des logs à 12 mois et du tracking pixel à 13 mois si non encore planifié ; "
           "arbitrer la durée de conservation des demandes DPO. [À COMPLÉTER]",
      size=9, italic=True, color=RED)

    out = "docs/rgpd/Referentiel-durees-conservation-BUUPP.docx"
    doc.save(out)
    return out


# ════════════════════════════════════════════════════════════════════
# DOC 3 — AIPD / DPIA
# ════════════════════════════════════════════════════════════════════
def build_aipd():
    doc = base_doc()
    cover(doc, "Analyse d'impact relative à la protection des données (AIPD)",
          "Traitement : Mise en relation rémunérée avec scoring et segmentation — art. 35 RGPD")
    doc.add_page_break()

    H(doc, "1. Faut-il une AIPD ? (critères CNIL / WP248)", 1)
    P(doc, "Une AIPD est requise lorsqu'un traitement est susceptible d'engendrer un risque "
           "élevé. Le traitement central de BUUPP remplit plusieurs des 9 critères du CEPD "
           "(WP248) — au moins deux suffisent à présumer le risque élevé :")
    BUL(doc, "Évaluation / scoring : calcul du BUUPP Score (qualité de profil).", "✔ ")
    BUL(doc, "Croisement / combinaison de données : segmentation multi-facettes (revenus, "
             "patrimoine, localisation, foyer, etc.).", "✔ ")
    BUL(doc, "Données à caractère hautement personnel : revenus, patrimoine, projets de vie, "
             "coordonnées bancaires (IBAN).", "✔ ")
    BUL(doc, "Usage innovant : modèle de rémunération de la donnée personnelle à double "
             "consentement (nouveauté).", "✔ ")
    BUL(doc, "Suivi systématique : journalisation des accès et des sollicitations.", "✔ ")
    P(doc, "Conclusion : AIPD requise pour ce traitement. Les traitements purement "
           "comptables/fiscaux (T7) et de paiement (T6, largement délégué à Stripe) "
           "présentent un risque résiduel plus faible mais alimentent ce traitement central ; "
           "ils sont couverts par la présente analyse. Les autres traitements (compte, "
           "communication, waitlist) ne nécessitent pas d'AIPD distincte mais sont documentés "
           "au registre.", bold=False)

    H(doc, "2. Description du traitement", 1)
    BUL(doc, "Mettre en relation des prospects consentants avec des professionnels, en "
             "valorisant et segmentant les profils, tout en rémunérant le prospect.", "Finalités : ")
    BUL(doc, "Identité (palier 1), profil enrichi (paliers 2-5 dont revenus/patrimoine), "
             "score dérivé, données de relation, coordonnées bancaires (rémunération).", "Données : ")
    BUL(doc, "Prospects (personnes concernées), professionnels (utilisateurs), personnels "
             "Majelink, sous-traitants (cf. registre).", "Acteurs : ")
    BUL(doc, "Collecte (inscription, consentement par palier) → scoring/segmentation → "
             "sollicitation (double consentement) → pseudonymisation → révélation journalisée "
             "→ rémunération → archivage/suppression selon le référentiel.", "Cycle de vie : ")
    BUL(doc, "Hébergement UE ; chiffrement ; RLS ; pseudonymisation à la lecture ; alias "
             "e-mail watermarqués ; journal des accès fail-closed + append-only ; anti-fraude.",
        "Mesures existantes : ")

    H(doc, "3. Nécessité et proportionnalité", 1)
    BUL(doc, "Exécution du contrat (mise en relation, paiement) ; consentement (paliers "
             "facultatifs, sollicitations) ; obligation légale (fiscal) ; intérêt légitime "
             "(fraude, sécurité). Légitimité à confirmer par la balance d'intérêts pour le "
             "scoring. [À COMPLÉTER]", "Base légale : ")
    BUL(doc, "Minimisation par finalité (seuls les paliers nécessaires à l'objectif du pro "
             "sont autorisés et facturés) ; paliers 2-5 facultatifs et révocables.", "Minimisation : ")
    BUL(doc, "Données saisies par le prospect lui-même ; téléphone vérifié par OTP ; "
             "unicité anti-doublon.", "Qualité / exactitude : ")
    BUL(doc, "Conformes au référentiel de conservation (document dédié).", "Durées : ")
    BUL(doc, "Information via la politique de confidentialité, la page DPO et la page "
             "minimisation ; double consentement explicite ; droits d'accès, rectification, "
             "effacement, opposition, portabilité exerçables.", "Information & droits : ")

    H(doc, "4. Appréciation des risques pour les personnes", 1)
    P(doc, "Trois événements redoutés (méthode CNIL) : accès illégitime, modification non "
           "désirée, disparition de données. Gravité et vraisemblance évaluées AVANT/APRÈS "
           "mesures. (Cotation à valider en atelier DPO — échelle : Négligeable / Limitée / "
           "Importante / Maximale.)", size=9, italic=True)
    table(doc,
          ["Risque", "Impact pour la personne", "Mesures en place", "Risque résiduel*"],
          [
            ["Accès illégitime aux données (ré-identification d'un profil, fuite d'IBAN/revenus)",
             "Atteinte à la vie privée, démarchage abusif, risque financier",
             "Pseudonymisation à la lecture ; alias watermarqués ; RLS ; chiffrement ; "
             "révélation gated + journalisée ; aucune donnée carte stockée",
             "Limité [à valider]"],
            ["Détournement de finalité par un pro (extraction, revente)",
             "Sollicitations hors cadre, perte de contrôle",
             "Données pseudonymisées non exportables ; consentement à usage unique ; "
             "watermark traçable ; journal des accès inviolable ; alerte abus",
             "Limité [à valider]"],
            ["Modification / falsification d'une trace d'accès",
             "Impossibilité de prouver qui a accédé aux données",
             "Journal append-only (trigger anti-UPDATE) + écriture fail-closed",
             "Négligeable [à valider]"],
            ["Profilage abusif (scoring) ou décision défavorable",
             "Exclusion d'opportunités, traitement inéquitable",
             "Score non décisionnel automatisé à effet juridique ; transparence ; "
             "pas de catégorie sensible ; droits d'opposition/rectification",
             "Limité [à valider]"],
            ["Disparition / indisponibilité des données",
             "Perte d'historique, de rémunération due",
             "Hébergement managé (sauvegardes Supabase) ; séquestre des fonds",
             "Limité [à valider]"],
          ], font=8)
    P(doc, "* La cotation finale (gravité × vraisemblance) doit être arrêtée en atelier avec "
           "le DPO et tracée dans l'outil PIA de la CNIL. [À COMPLÉTER]", size=8.5, italic=True, color=RED)

    H(doc, "5. Plan d'action et validation", 1)
    BUL(doc, "Formaliser la balance d'intérêts pour le scoring (intérêt légitime).")
    BUL(doc, "Arbitrer et automatiser les purges manquantes (prospection 3 ans, logs 12 mois).")
    BUL(doc, "Tenir à jour les DPA des sous-traitants ; harmoniser la région d'hébergement "
             "annoncée (Irlande) avec la politique publiée.")
    BUL(doc, "Réaliser/valider la cotation des risques dans l'outil PIA CNIL avec le DPO.")
    BUL(doc, "Avis du DPO et validation par le responsable de traitement à consigner ci-dessous.")
    P(doc, "")
    table(doc, ["Rôle", "Nom", "Avis / décision", "Date"],
          [["Responsable de traitement", "[À COMPLÉTER]", "", ""],
           ["DPO", "[À COMPLÉTER]", "", ""],
           ["Avis des personnes concernées (le cas échéant)", "[À COMPLÉTER]", "", ""]],
          font=9)

    P(doc, "Cette AIPD est un document de travail à valider par le DPO. Elle ne constitue "
           "pas un avis juridique. Outil recommandé : logiciel PIA de la CNIL.",
      italic=True, size=8.5, color=GREY)

    out = "docs/rgpd/AIPD-BUUPP-mise-en-relation-scoring.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    for f in (build_registre(), build_referentiel(), build_aipd()):
        print("OK ->", f)
