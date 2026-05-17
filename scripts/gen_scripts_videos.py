#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère docs/scripts-videos-buupp.docx — 17 scripts vidéo (8 Pro + 9 Prospect)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import os

INK = RGBColor(0x0F, 0x16, 0x29)
ACCENT = RGBColor(0x6D, 0x5B, 0xFF)
AMBER = RGBColor(0xB4, 0x53, 0x09)
GREY = RGBColor(0x5B, 0x64, 0x78)

doc = Document()

# ---- styles de base ----
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(11)
base.font.color.rgb = INK


def H(txt, size=20, color=INK, after=6, before=14, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    return p


def label(p_label, text, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(p_label + " ")
    r.bold = True
    r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(3)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def seq(time, vo, visual):
    p = doc.add_paragraph()
    rt = p.add_run(f"[{time}]  ")
    rt.bold = True
    rt.font.color.rgb = ACCENT
    rt.font.size = Pt(10.5)
    rv = p.add_run("VOIX-OFF — ")
    rv.bold = True
    rv.font.size = Pt(10.5)
    rb = p.add_run(f"« {vo} »")
    rb.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    pv = doc.add_paragraph()
    pv.paragraph_format.left_indent = Inches(0.35)
    rvi = pv.add_run(visual)
    rvi.italic = True
    rvi.font.size = Pt(10)
    rvi.font.color.rgb = GREY
    pv.paragraph_format.space_after = Pt(8)


def script(num, space, tab, dur, hook_vo, hook_screen, sequences, cta, onscreen, tags):
    H(f"SCRIPT {num} — Espace {space} · Onglet « {tab} »", size=16, color=ACCENT,
      before=20, after=4)
    meta = doc.add_paragraph()
    rm = meta.add_run(f"Durée cible : {dur}  ·  Format : vertical 9:16 (TikTok/Reels) "
                      f"+ 1:1/16:9 (LinkedIn/Facebook)  ·  Voix-off + sous-titres incrustés")
    rm.font.size = Pt(9.5)
    rm.font.color.rgb = GREY
    meta.paragraph_format.space_after = Pt(8)

    label("🎬 ACCROCHE (0–3 s) —", f"« {hook_vo} »", color=AMBER)
    pe = doc.add_paragraph()
    pe.paragraph_format.left_indent = Inches(0.35)
    re = pe.add_run(f"Texte plein écran : “{hook_screen}”. Plan : capture de l'onglet "
                    f"qui apparaît en fondu/zoom rapide.")
    re.italic = True
    re.font.size = Pt(10)
    re.font.color.rgb = GREY
    pe.paragraph_format.space_after = Pt(8)

    H("Déroulé :", size=11, color=INK, before=4, after=4)
    for t, vo, vis in sequences:
        seq(t, vo, vis)

    label("📣 CALL TO ACTION (fin) —", f"« {cta} »", color=ACCENT)
    H("Textes à incruster (sous-titres clés) :", size=10.5, color=INK, before=6, after=2)
    for o in onscreen:
        bullet(o)
    pt = doc.add_paragraph()
    rt = pt.add_run("Hashtags : ")
    rt.bold = True
    rt.font.size = Pt(9.5)
    rt2 = pt.add_run(tags)
    rt2.font.size = Pt(9.5)
    rt2.font.color.rgb = GREY
    pt.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("— — —").alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================ PAGE DE GARDE ============================
H("BUUPP", size=34, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, before=60, after=2)
H("17 scripts vidéo — Espaces Pro & Prospect", size=18, color=INK,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Réseaux sociaux : LinkedIn · Facebook · TikTok · Instagram")
rs.font.color.rgb = GREY
rs.font.size = Pt(11)

doc.add_paragraph()
H("Comment utiliser ce document", size=13, color=INK, before=18, after=4)
for b in [
    "8 scripts pour l'espace Pro, 9 pour l'espace Prospect (17 vidéos courtes, 1–2 min).",
    "Chaque script = un onglet : accroche (0–3 s), déroulé minuté, call to action.",
    "Les didascalies en italique gris décrivent le plan à filmer/capturer et "
    "l'endroit où placer la flèche/zoom au montage.",
    "🎬 = ce qu'on montre à l'écran.  🔍 = élément à zoomer + flèche animée dessus.",
    "Toutes les VOIX-OFF sont prêtes à être lues telles quelles (langage grand public, "
    "zéro jargon).",
]:
    bullet(b)

note = doc.add_paragraph()
rn = note.add_run("Note de livraison : ")
rn.bold = True
rn.font.color.rgb = AMBER
rn2 = note.add_run(
    "ce fichier a été généré localement (impossible d'envoyer un e-mail "
    "automatiquement). Pense à l'envoyer toi-même à jjlex64@gmail.com.")
rn2.font.color.rgb = INK
note.paragraph_format.space_before = Pt(14)

doc.add_page_break()

# ============================ ESPACE PRO ============================
H("ESPACE PRO — 8 scripts", size=22, color=ACCENT, before=4, after=2)
doc.add_paragraph(
    "Le pro lance des campagnes de mise en relation et paie des particuliers "
    "qui ont accepté d'être contactés.").runs[0].font.color.rgb = GREY

script(
    1, "Pro", "Créer une campagne", "≈ 90 s",
    "Et si vous choisissiez exactement QUI vous contacte — et qu'eux soient payés pour accepter ?",
    "Lancez une campagne en 8 étapes",
    [
        ("3–12 s",
         "Avec BUUPP, vous ne payez plus pour du clic au hasard. Vous créez une campagne guidée, étape par étape.",
         "🎬 Plan large sur le wizard. 🔍 Zoom + flèche sur la barre de progression : Objectif · Dates · Données · Ciblage · Budget · Mots-clés · Description · Récap."),
        ("12–30 s",
         "D'abord, votre objectif : un appel, un rendez-vous, un événement, un téléchargement… Puis votre cible : ville, département, région ou national, tranche d'âge, niveau de vérification.",
         "🎬 Étape Objectif (grille de catégories) puis étape Ciblage. 🔍 Flèche sur les 4 cases Ville / Département / Région / National."),
        ("30–48 s",
         "Vous choisissez les données dont vous avez besoin — uniquement celles-là. C'est le RGPD, intégré par design : le principe de minimisation, article 5 du RGPD.",
         "🎬 Étape Données : les 5 paliers. 🔍 Zoom + flèche sur le bandeau « Principe de minimisation — RGPD »."),
        ("48–68 s",
         "Et surtout : le budget. Vous voyez en temps réel le coût par contact, la commission, et le total exact avant de valider. Aucune mauvaise surprise.",
         "🔍 GROS PLAN avec flèche sur le slider « Contacts souhaités » et l'encart live « Total à débiter »."),
        ("68–82 s",
         "Un dernier récapitulatif clair, vous validez les conditions, et vous lancez.",
         "🎬 Étape Récap (tableau de synthèse). 🔍 Flèche animée qui pointe le bouton « Lancer la campagne »."),
    ],
    "Votre première campagne ciblée et conforme RGPD vous attend sur BUUPP.",
    ["Vous choisissez qui vous contacte", "Le RGPD, intégré par design",
     "Coût exact AVANT de lancer", "8 étapes, zéro jargon"],
    "#BUUPP #marketing #RGPD #prospection #leadgen",
)

script(
    2, "Pro", "Vue d'ensemble", "≈ 70 s",
    "Vos campagnes sont-elles rentables ? Réponse en un coup d'œil.",
    "Votre tableau de bord, en 3 secondes",
    [
        ("3–14 s",
         "Dès la connexion, l'essentiel est là : votre crédit disponible, vos contacts du mois, vos campagnes actives.",
         "🎬 Plan sur l'en-tête Pro. 🔍 Zoom + flèche sur le crédit disponible et « N contacts ce mois »."),
        ("14–34 s",
         "Quatre indicateurs clés : contacts acceptés sur 30 jours, taux d'acceptation, coût moyen par contact, et votre retour sur investissement estimé.",
         "🎬 Les 4 cartes KPI. 🔍 Flèche qui balaie les 4 cartes une à une."),
        ("34–52 s",
         "Vous vous demandez d'où vient ce ROI ? Un clic, et BUUPP vous explique le calcul avec VOS chiffres réels.",
         "🔍 GROS PLAN + flèche sur la petite icône « i » de la carte ROI, puis ouverture de la modale « Comment on calcule votre ROI ? »."),
        ("52–66 s",
         "En dessous : la performance de vos campagnes sur 7, 30 ou 90 jours, et vos dernières acceptations en direct.",
         "🎬 Histogramme de performance + tableau « Dernières acceptations ». 🔍 Flèche sur les chips 7J / 30J / 90J."),
    ],
    "Pilotez vos campagnes au chiffre près, pas au feeling. C'est BUUPP.",
    ["Crédit, contacts, ROI : tout de suite", "Le ROI expliqué avec VOS chiffres",
     "Performance sur 7 / 30 / 90 j"],
    "#BUUPP #ROI #dashboard #marketingdigital",
)

script(
    3, "Pro", "Campagnes", "≈ 70 s",
    "Toutes vos campagnes. Un seul écran. Zéro prise de tête.",
    "Gérez tout depuis un seul endroit",
    [
        ("3–16 s",
         "Vos initiatives en cours, filtrées en un clic : actives, en pause, terminées.",
         "🎬 Liste des campagnes. 🔍 Flèche sur les filtres « Toutes / Actives / En pause / Terminées »."),
        ("16–34 s",
         "Pour chaque campagne : budget consommé, prospects touchés, contacts obtenus. Tout est visible.",
         "🎬 Une carte campagne. 🔍 Zoom + flèche sur les 3 compteurs Budget / Touchés / Contacts et la barre de progression."),
        ("34–52 s",
         "Et ce code, c'est votre sésame : le Code BUUPP, à communiquer au prospect lors du contact. Une preuve, anti-fraude.",
         "🔍 GROS PLAN + flèche animée sur le badge orange « 🔒 Code BUUPP XXXX »."),
        ("52–66 s",
         "Besoin d'une pause ? 48 heures chrono, sans perdre vos acceptations. Vous pouvez aussi dupliquer une campagne qui marche.",
         "🎬 Clic sur « Pause » → modale « Pause café · 48 h chrono ». 🔍 Flèche sur les boutons « Pause » et « Dupliquer »."),
    ],
    "Vos campagnes sous contrôle, à tout moment. BUUPP.",
    ["Filtrez en un clic", "Le Code BUUPP = preuve anti-fraude",
     "Pause 48 h sans rien perdre"],
    "#BUUPP #campagnes #gestion #antifraude",
)

script(
    4, "Pro", "Mes contacts", "≈ 75 s",
    "Des contacts qui ont dit OUI. Pas des numéros achetés au hasard.",
    "Vos contacts, vraiment consentants",
    [
        ("3–16 s",
         "Ici, tous les prospects qui ont accepté votre campagne, regroupés par campagne.",
         "🎬 Carnet de contacts par campagne. 🔍 Flèche sur l'en-tête d'un groupe campagne."),
        ("16–34 s",
         "Filtrez par BUUPP Score, par contact atteint, par palier de données. Vous parlez aux bonnes personnes, point.",
         "🎬 Les pills de filtres combinés. 🔍 Zoom + flèche sur « Score ≥ 720 » et « Palier 2 »."),
        ("34–52 s",
         "Contactez-les par le canal qu'ILS ont partagé : e-mail, appel, SMS, WhatsApp. Et un message groupé qui protège chaque adresse — RGPD oblige.",
         "🔍 GROS PLAN + flèche sur le bouton « Message groupé (N) », puis aperçu de l'envoi en copie cachée."),
        ("52–70 s",
         "Chaque fiche est filigranée. Toute fuite hors campagne déclenche une enquête automatique. Vos données et les leurs sont protégées.",
         "🎬 Bandeau de politique d'usage en bas. 🔍 Flèche sur la mention « watermarking appliqué à chaque fiche »."),
    ],
    "Des leads consentants, contactables, et protégés. C'est ça, BUUPP.",
    ["Ils ont accepté d'être contactés", "Filtrez par score et palier",
     "Message groupé protégé RGPD"],
    "#BUUPP #leadgen #consentement #RGPD",
)

script(
    5, "Pro", "Analytics", "≈ 70 s",
    "À quelle heure vos prospects disent OUI ? BUUPP le sait.",
    "Vos performances, à la loupe",
    [
        ("3–16 s",
         "Filtrez par campagne et par période, et voyez ce qui marche vraiment.",
         "🎬 Les deux sélecteurs Campagne / Période. 🔍 Flèche sur « 30 derniers jours »."),
        ("16–32 s",
         "Quels paliers de données convertissent le mieux ? La réponse est en barres, claire.",
         "🎬 Carte « Taux d'acceptation par palier ». 🔍 Flèche qui descend les paliers P1 → P5."),
        ("32–52 s",
         "Et voici la pépite : la carte des meilleurs créneaux. Les jours et les heures où vos prospects acceptent le plus.",
         "🔍 GROS PLAN + flèche sur la heatmap « Meilleurs créneaux » et le résumé « Top créneaux · Lundi 14h… »."),
        ("52–66 s",
         "D'où viennent vos contacts, quel âge, quel profil : tout est segmenté pour affiner la prochaine campagne.",
         "🎬 Cartes Géographie, Tranche d'âge, Sexe. 🔍 Flèche balayant les 3 cartes."),
    ],
    "Arrêtez de deviner. Lancez au bon moment, sur la bonne cible. BUUPP.",
    ["Quel palier convertit le mieux", "Le meilleur créneau pour lancer",
     "Géo, âge, profil : tout segmenté"],
    "#BUUPP #analytics #data #marketing",
)

script(
    6, "Pro", "Mes informations", "≈ 70 s",
    "Un profil pro vérifié, des factures conformes — sans paperasse.",
    "Votre identité pro, carrée",
    [
        ("3–16 s",
         "Renseignez une fois votre société : raison sociale, forme juridique, adresse. BUUPP vérifie votre activité.",
         "🎬 Grille « Informations société ». 🔍 Flèche sur la carte « Complétude du profil entreprise · N/4 »."),
        ("16–32 s",
         "Votre SIREN ? Strictement confidentiel, utilisé seulement pour vérifier votre existence légale.",
         "🔍 Zoom + flèche sur la bannière « Votre SIREN reste strictement confidentiel »."),
        ("32–54 s",
         "Et ici, vous changez de formule en un clic : Starter ou Pro. Passer en Pro débloque plus de paliers et jusqu'à 500 prospects par campagne.",
         "🔍 GROS PLAN + flèche animée sur le sélecteur de plan, bouton « Passer en Pro »."),
        ("54–66 s",
         "Tout s'enregistre automatiquement, et les numéros officiels sont vérifiés en direct sur les registres publics.",
         "🎬 Modale d'édition. 🔍 Flèche sur « Modifications enregistrées automatiquement » et le statut « Validé »."),
    ],
    "Profil conforme, abonnement maîtrisé. Tout au même endroit. BUUPP.",
    ["SIREN confidentiel", "Starter ou Pro en un clic",
     "Vérification officielle en direct"],
    "#BUUPP #entreprise #conformité #SaaS",
)

script(
    7, "Pro", "Facturation", "≈ 65 s",
    "Vos factures, prêtes pour le comptable. Vraiment prêtes.",
    "Paiements & factures, zéro stress",
    [
        ("3–16 s",
         "En haut : votre abonnement, la date de renouvellement, la carte enregistrée. Clair.",
         "🎬 Les 3 cartes de synthèse. 🔍 Flèche sur « Abonnement actuel » et « Renouvellement »."),
        ("16–34 s",
         "En dessous : l'historique complet de vos factures, avec leur statut — payée, en attente, échouée.",
         "🎬 Tableau « Historique des factures ». 🔍 Flèche sur les chips de statut « ✓ Payée »."),
        ("34–54 s",
         "Vous téléchargez une facture ? BUUPP vérifie votre SIREN sur le registre officiel et complète les mentions légales avant de générer le PDF. Conforme du premier coup.",
         "🔍 GROS PLAN + flèche sur le bouton « PDF », puis la modale « Compléter les mentions légales » avec le statut « Validé »."),
    ],
    "Des factures conformes, en deux clics. C'est BUUPP.",
    ["Abo, renouvellement, carte : visibles", "SIREN vérifié officiellement",
     "PDF conforme du premier coup"],
    "#BUUPP #facturation #comptabilité #conformité",
)

script(
    8, "Pro", "Mes messages", "≈ 55 s",
    "Les annonces importantes de BUUPP : jamais ratées.",
    "Votre boîte de réception BUUPP",
    [
        ("3–16 s",
         "Toutes les communications officielles de l'équipe BUUPP, centralisées ici. Pas dans un e-mail perdu.",
         "🎬 Liste des messages. 🔍 Flèche sur le sous-titre « N message(s) · N non lu(s) »."),
        ("16–34 s",
         "Un point de couleur, un badge : vous voyez tout de suite ce que vous n'avez pas lu. Un clic, le message s'ouvre et se marque comme lu.",
         "🔍 GROS PLAN + flèche sur le badge « Non lu », puis le dépliage de la carte message."),
        ("34–48 s",
         "Une pièce jointe ? Elle se télécharge directement, sans quitter votre tableau de bord.",
         "🎬 Message déplié avec pièce jointe. 🔍 Flèche sur « ⬇ Télécharger »."),
    ],
    "Restez informé, sans rien chercher. BUUPP.",
    ["Communications officielles centralisées", "Non lu repéré en un clin d'œil",
     "Pièces jointes téléchargeables"],
    "#BUUPP #notifications #produit",
)

doc.add_page_break()

# ============================ ESPACE PROSPECT ============================
H("ESPACE PROSPECT — 9 scripts", size=22, color=ACCENT, before=4, after=2)
doc.add_paragraph(
    "Le prospect (un particulier) est payé pour accepter d'être contacté, "
    "et garde le contrôle total de ses données.").runs[0].font.color.rgb = GREY

script(
    1, "Prospect", "Portefeuille", "≈ 75 s",
    "Vos données vous ont rapporté combien ce mois-ci ? Regardez.",
    "Vos gains, en vrai",
    [
        ("3–16 s",
         "Vos gains du mois s'affichent dès l'arrivée. Sur BUUPP, vos données, c'est vous qui les monétisez.",
         "🎬 Bandeau d'en-tête. 🔍 Flèche sur « Vos gains du mois : X € »."),
        ("16–36 s",
         "Trois soldes, ultra clairs : ce qui est disponible maintenant, ce qui est en séquestre en attendant la fin d'une campagne, et tout ce que vous avez cumulé depuis le début.",
         "🎬 Les 3 cartes de solde. 🔍 Flèche sur « Disponible », « En séquestre », « Cumulé depuis ouverture »."),
        ("36–58 s",
         "Et le moment qu'on préfère : vous retirez vos gains. Un virement sur votre compte, sous quelques jours.",
         "🔍 GROS PLAN + flèche animée sur le bouton « Retirer mes gains → », puis la modale de retrait par virement."),
        ("58–70 s",
         "Chaque euro est tracé dans l'historique des mouvements : date, origine, palier, montant.",
         "🎬 Tableau « Historique des mouvements ». 🔍 Flèche balayant les colonnes."),
    ],
    "Vos données ont de la valeur. Récupérez-la. Rejoignez BUUPP.",
    ["Vos gains du mois, tout de suite", "Disponible / Séquestre / Cumulé",
     "Retrait par virement bancaire"],
    "#BUUPP #revenucomplémentaire #tesdonnées #cashback",
)

script(
    2, "Prospect", "Mes données", "≈ 85 s",
    "Vos données. Vos règles. Pour de vrai.",
    "Le contrôle total, c'est ici",
    [
        ("3–14 s",
         "Sur BUUPP, vos données vous appartiennent. Et la loi vous protège : RGPD, articles 12 à 22.",
         "🎬 Onglet Mes données. 🔍 Flèche sur le bandeau jaune « Vos droits sur vos données — RGPD »."),
        ("14–34 s",
         "Elles sont rangées en 5 paliers : identification, localisation, style de vie, données pro, patrimoine. Vous voyez votre progression en temps réel.",
         "🎬 Carte de progression. 🔍 Zoom + flèche sur les 5 barres P1 → P5."),
        ("34–56 s",
         "Pour chaque donnée, vous décidez : ajouter, modifier, masquer temporairement… ou supprimer définitivement. À tout instant.",
         "🔍 GROS PLAN + flèche sur les boutons « Ajouter », « Masquer temporairement », « Supprimer »."),
        ("56–76 s",
         "Et plus vous remplissez, plus vous devenez éligible aux campagnes les mieux payées. Vous gardez la main, vous maximisez vos gains.",
         "🎬 Bandeau de nudge « Encore un petit effort… 😉 ». 🔍 Flèche sur un palier presque complété."),
    ],
    "Reprenez le pouvoir sur vos données — et soyez payé pour. BUUPP.",
    ["RGPD : vos droits, articles 12 à 22", "5 paliers, progression en direct",
     "Ajouter, masquer, supprimer : vous décidez"],
    "#BUUPP #RGPD #datprivacy #tesdonnées",
)

script(
    3, "Prospect", "Mises en relation", "≈ 80 s",
    "Une entreprise veut vous parler. C'est VOUS qui décidez.",
    "Vous acceptez. Ou pas.",
    [
        ("3–16 s",
         "Une marque vous sollicite ? Vous voyez qui, pourquoi, et combien ça vous rapporte. Avant de répondre.",
         "🎬 Grille des demandes en attente. 🔍 Flèche sur le nom du pro et le secteur."),
        ("16–36 s",
         "Le montant est affiché en grand. Un compte à rebours indique le temps qu'il vous reste pour décider.",
         "🔍 GROS PLAN + flèche sur la récompense (ex. « 2,50 € ») et le timer « ⚡ 14 h 22 min »."),
        ("36–56 s",
         "Vous acceptez ? Votre récompense passe en séquestre et vous est créditée après le contact. Et c'est réversible tant que la campagne n'est pas close.",
         "🔍 Flèche animée sur le bouton « Accepter », puis le bloc vert « Paiement en séquestre »."),
        ("56–74 s",
         "Votre accord vaut pour cette campagne uniquement. Pas de revente, pas de réutilisation. Et tout votre historique est là, filtrable.",
         "🎬 Bandeau « Accord strictement limité à cette campagne ». 🔍 Flèche sur les filtres « Acceptées / Refusées » et un badge « ⚡ FLASH »."),
    ],
    "Être contacté, oui — mais à vos conditions, et payé. BUUPP.",
    ["Qui, pourquoi, combien : avant de dire oui", "Paiement sécurisé en séquestre",
     "Accord limité à une seule campagne"],
    "#BUUPP #consentement #revenucomplémentaire",
)

script(
    4, "Prospect", "Paliers de vérification", "≈ 70 s",
    "Plus vous êtes vérifié, plus vous gagnez. Logique.",
    "Montez en niveau de confiance",
    [
        ("3–18 s",
         "Trois paliers de confiance : Basique à la création, Vérifié, puis Certifié confiance. Chaque palier débloque des demandes plus exigeantes — et mieux payées.",
         "🎬 Barre de progression à 3 pastilles. 🔍 Flèche sur « Palier actuel » et « Prochaine étape »."),
        ("18–40 s",
         "Pour passer Vérifié : vérifiez votre téléphone, et renseignez votre RIB. Deux minutes, et vous accédez à de nouvelles opportunités.",
         "🔍 GROS PLAN + flèche animée sur le bouton « Renseigner mon RIB → »."),
        ("40–58 s",
         "Le niveau Certifié confiance ? Il s'obtient en acceptant un rendez-vous physique avec un professionnel. Le top de la rémunération.",
         "🎬 Carte « Certifié confiance ». 🔍 Flèche sur le chip « À venir » / « Validé ✓ »."),
    ],
    "Vérifiez-vous, débloquez les meilleures offres. BUUPP.",
    ["3 paliers : Basique, Vérifié, Certifié", "RIB = palier Vérifié débloqué",
     "Plus de confiance = mieux payé"],
    "#BUUPP #confiance #revenucomplémentaire",
)

script(
    5, "Prospect", "BUUPP Score", "≈ 75 s",
    "Votre score sur 1000. C'est votre valeur aux yeux des marques.",
    "Découverte → Prestige",
    [
        ("3–16 s",
         "Le BUUPP Score, c'est votre indice de désirabilité, sur 1000 points.",
         "🔍 GROS PLAN + flèche sur la jauge circulaire et son niveau : Découverte, Solide, Recherchée, Prestige."),
        ("16–36 s",
         "Il se calcule sur trois choses : la complétude de vos paliers, la fraîcheur de vos données, et votre taux d'acceptation.",
         "🎬 Les 3 mini-barres sous la jauge. 🔍 Flèche sur chacune des 3 dimensions."),
        ("36–56 s",
         "Vous voyez son évolution dans le temps, et surtout : des conseils chiffrés pour le faire grimper. « Encore X points pour le niveau supérieur. »",
         "🎬 Courbe d'évolution + bandeau conseils. 🔍 Flèche sur « Encore X pts pour atteindre [niveau] »."),
        ("56–70 s",
         "Un score plus haut, c'est plus de sollicitations, et mieux rémunérées. Tout est expliqué, rien n'est caché.",
         "🎬 Les 3 cartes conseils. 🔍 Flèche sur un gain potentiel « +N pts max »."),
    ],
    "Faites grimper votre score, faites grimper vos gains. BUUPP.",
    ["Un score de désirabilité sur 1000", "3 leviers clairs pour progresser",
     "Conseils chiffrés, zéro boîte noire"],
    "#BUUPP #score #gamification #revenu",
)

script(
    6, "Prospect", "Préférences", "≈ 75 s",
    "Qui peut vous contacter ? C'est vous qui tracez la ligne.",
    "Vous filtrez. Eux s'adaptent.",
    [
        ("3–18 s",
         "Choisissez les types de campagnes que vous acceptez : un appel, un rendez-vous, un événement, une promo… ou tout.",
         "🎬 Carte « Types de campagne acceptés ». 🔍 Flèche sur les pilules de types."),
        ("18–36 s",
         "Choisissez aussi les secteurs : immobilier, finance, beauté, auto… Onze catégories, à vous de cocher.",
         "🎬 Carte « Catégories autorisées ». 🔍 Flèche balayant les pilules de secteurs."),
        ("36–56 s",
         "Définissez votre zone : un rayon autour de chez vous, ou la France entière. Le curseur, c'est vous.",
         "🔍 GROS PLAN + flèche sur le slider de rayon et la case « Étendre au niveau national »."),
        ("56–72 s",
         "Et le plus important : quels paliers de données vous acceptez de partager, avec la fourchette de gains pour chacun. Vous arbitrez vie privée et revenus.",
         "🔍 Flèche animée sur la liste « Paliers partageables » et les fourchettes « 5,00 – 10,00 € »."),
    ],
    "Vos données, votre périmètre, vos conditions. BUUPP.",
    ["Types & secteurs : vous cochez", "Rayon local ou national",
     "Vous arbitrez vie privée / revenus"],
    "#BUUPP #contrôle #datprivacy #consentement",
)

script(
    7, "Prospect", "Parrainage", "≈ 65 s",
    "Invitez vos proches. Gagnez plus. Tout le monde y gagne.",
    "Votre lien = des bonus",
    [
        ("3–16 s",
         "Vous avez un lien de parrainage unique. Vous le copiez, vous le partagez, c'est parti.",
         "🔍 GROS PLAN + flèche sur « buupp.com/ref/[CODE] » et le bouton « Copier » (feedback « Copié ! »)."),
        ("16–36 s",
         "Suivez vos filleuls en direct : combien sont actifs, combien de places restantes, votre bonus actuel.",
         "🎬 Les 4 cartes statistiques. 🔍 Flèche sur « Filleuls actifs » et « Bonus actuel »."),
        ("36–54 s",
         "À 10 filleuls, vous débloquez le statut VIP : un bonus exceptionnel sur les grosses campagnes. Et vos filleuls profitent aussi du bonus de lancement.",
         "🎬 Bannière dorée « Palier VIP atteint ». 🔍 Flèche sur le statut « VIP · Palier débloqué »."),
    ],
    "Partagez BUUPP, multipliez vos gains. Rejoignez le mouvement.",
    ["Un lien unique, un clic pour copier", "Filleuls suivis en direct",
     "10 filleuls = statut VIP"],
    "#BUUPP #parrainage #VIP #communauté",
)

script(
    8, "Prospect", "Informations fiscales", "≈ 70 s",
    "Zéro mauvaise surprise aux impôts. BUUPP s'en occupe.",
    "Vos seuils, sous contrôle",
    [
        ("3–18 s",
         "Tous vos gains de l'année sont récapitulés, en toute transparence.",
         "🎬 Carte « Exercice [année] (en cours) ». 🔍 Flèche sur le total annuel."),
        ("18–40 s",
         "Une barre vous montre où vous en êtes par rapport au seuil de déclaration : 2 000 € ou 30 transactions. Vous savez exactement quand ça compte.",
         "🔍 GROS PLAN + flèche sur la barre « Seuil déclaratif : X / 2 000 € »."),
        ("40–58 s",
         "Trois repères clés expliqués simplement : 305 €, 2 000 €, 77 700 €. Pas besoin d'être expert.",
         "🎬 Carte « Seuils à retenir ». 🔍 Flèche sur les 3 montants."),
        ("58–68 s",
         "Et vos documents — récap annuel, reçu DGFiP — se téléchargent en un clic.",
         "🔍 Flèche sur les boutons « Récap [année] (PDF) » et « Reçu DGFiP »."),
    ],
    "Gagnez l'esprit tranquille. BUUPP gère la transparence fiscale.",
    ["Gains annuels récapitulés", "Seuil DGFiP suivi en direct",
     "Documents fiscaux téléchargeables"],
    "#BUUPP #fiscalité #transparence #DAC7",
)

script(
    9, "Prospect", "Mes messages", "≈ 55 s",
    "Restez informé. Sans jamais avoir à chercher.",
    "Les annonces BUUPP, ici",
    [
        ("3–16 s",
         "Toutes les annonces de l'équipe BUUPP arrivent directement dans votre espace. Pas dans un mail noyé.",
         "🎬 Liste des messages. 🔍 Flèche sur « X messages · Y non lus »."),
        ("16–34 s",
         "Un badge « Non lu », et vous savez tout de suite s'il y a du nouveau : un changement de conditions, un nouveau palier…",
         "🔍 GROS PLAN + flèche sur le badge « Non lu », puis l'ouverture du message."),
        ("34–48 s",
         "Un document à récupérer ? Il se télécharge sur place. Simple.",
         "🎬 Message avec pièce jointe. 🔍 Flèche sur « Télécharger [fichier] »."),
    ],
    "L'info qui compte, là où il faut. BUUPP.",
    ["Annonces centralisées dans l'espace", "Non lu repéré instantanément",
     "Pièces jointes à portée de clic"],
    "#BUUPP #produit #communauté",
)

# ---- conseils de tournage (annexe) ----
doc.add_page_break()
H("Annexe — Conseils de tournage & montage", size=18, color=ACCENT, after=6)
for b in [
    "Capturez l'écran réel de chaque onglet (démo connectée) en 1080p minimum ; "
    "le format vertical 9:16 cadre serré sur l'élément clé du « 🔍 GROS PLAN ».",
    "Flèche/zoom : utilisez une flèche animée qui « pousse » vers l'élément (effet de "
    "scale + léger surlignage), synchronisée avec la phrase de voix-off correspondante.",
    "Rythme : coupez toutes les 2–4 s, sous-titres incrustés en gros (lecture sans son).",
    "Accroche : la phrase « 0–3 s » doit être prononcée ET affichée en plein écran, "
    "avant tout logo, pour retenir le scroll.",
    "CTA de fin : carton BUUPP + texte « Crée ton compte sur buupp.com » 1,5 s.",
    "Cohérence : même charte (police, couleur d'accent violet, même voix) sur les 17 vidéos "
    "pour une série reconnaissable.",
]:
    bullet(b)

# ---- save ----
out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
os.makedirs(out_dir, exist_ok=True)
out = os.path.join(out_dir, "scripts-videos-buupp.docx")
doc.save(out)
print("OK ->", out)
