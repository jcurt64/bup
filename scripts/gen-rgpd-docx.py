#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère le document Word de conformité RGPD/sécurité de BUUPP, destiné à être
présenté lors d'un contrôle (CNIL ou autre autorité).

Usage : python3 scripts/gen-rgpd-docx.py
Sortie : docs/rgpd/BUUPP-conformite-securite-RGPD.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
INK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x47, 0x55, 0x65)

doc = Document()

# Styles de base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = INDIGO if level <= 1 else INK
    return p

def para(text, italic=False, bold=False, color=None, size=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def proof(text):
    p = doc.add_paragraph()
    r = p.add_run("Preuve / vérifiable dans : ")
    r.bold = True
    r.font.color.rgb = GREY
    r.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.color.rgb = GREY
    r2.font.size = Pt(9)
    r2.font.name = "Consolas"
    p.paragraph_format.space_after = Pt(10)
    return p

def evidence_table(rows):
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for c, label in zip(hdr, ["Mesure", "Mise en œuvre", "Emplacement (preuve)"]):
        c.paragraphs[0].add_run(label).bold = True
    for r in rows:
        cells = t.add_row().cells
        for c, val in zip(cells, r):
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            if c is cells[2]:
                run.font.name = "Consolas"
    return t

# ── Page de garde ────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("BUUPP")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = INDIGO

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Mesures de sécurité et de protection des données personnelles")
r.font.size = Pt(15)
r.font.color.rgb = INK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Dossier de conformité RGPD — à présenter en cas de contrôle (CNIL ou autre autorité)")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Éditeur : Majelink — Plateforme BUUPP\n").bold = True
meta.add_run("RCS Pau 892 514 167 · 12 Impasse des Étriers, 64140 Lons\n")
meta.add_run("Contact DPO : dp.buupp@buupp.com · Contact : contact@buupp.com\n")
meta.add_run("Version du document : 14/06/2026")

doc.add_page_break()

# ── 0. Préambule ─────────────────────────────────────────────────────
h("0. Objet du document", 1)
para("BUUPP est une plateforme française de mise en relation rémunérée à double "
     "consentement entre des particuliers (« prospects ») et des professionnels. "
     "Les particuliers acceptent, sollicitation par sollicitation, d'être "
     "contactés en échange d'une rémunération ; aucune donnée n'est transmise "
     "sans leur accord explicite.")
para("Ce document recense les mesures techniques et organisationnelles mises en "
     "œuvre pour assurer la sécurité et la protection des données personnelles, "
     "et indique pour chacune où la vérifier dans le système (fichier de code, "
     "migration de base de données, configuration). Il est destiné à répondre à "
     "un contrôle de la CNIL ou d'une autre autorité.")
para("Principe directeur : la protection des données est intégrée dès la "
     "conception (privacy by design) et par défaut (privacy by default). "
     "L'expertise RGPD est portée par le fondateur, juriste spécialisé en "
     "protection des données personnelles.", italic=True)

# ── 1. Pseudonymisation & minimisation ───────────────────────────────
h("1. Pseudonymisation et minimisation des données", 1)
para("Les données des prospects sont stockées en base mais ne sont JAMAIS "
     "transmises telles quelles au professionnel. La transformation est "
     "appliquée à la lecture, côté serveur :")
bullet("nom → masquage (initiale + points) ; e-mail → alias watermarqué ; "
       "date de naissance → tranche d'âge de 5 ans ; code postal → département ; "
       "adresse précise / poste / revenus → suppression (jamais transmis) ; "
       "véhicule, animaux, foyer → catégorisation.")
bullet("Conséquence : seul BUUPP peut relier un profil pseudonymisé à la "
       "personne réelle (« réversible par buupp seul »).")
proof("lib/pro/pseudonymize.ts ; appliqué dans app/api/pro/contacts/[relationId]/details/route.ts")
para("Minimisation par finalité : lorsqu'un professionnel crée une campagne, "
     "seuls les paliers de données strictement nécessaires à l'objectif choisi "
     "sont autorisés et facturés. Rien de superflu ne circule.")
proof("Wizard de campagne (Objectives.jsx / Pro.jsx) ; page publique /minimisation")

# ── 2. Consentement ──────────────────────────────────────────────────
h("2. Consentement (double consentement)", 1)
para("Deux consentements explicites sont requis avant toute transmission : "
     "(1) l'acceptation des CGU par le prospect à l'inscription, (2) l'acceptation "
     "explicite de chaque sollicitation d'un professionnel. Chaque accord est "
     "strictement limité à une seule sollicitation — pas de réutilisation ni de "
     "revente.")
proof("lib/cnil/consent.ts ; flux d'acceptation des sollicitations ; CGU /cgu, /cgv")

# ── 3. Sécurité technique ────────────────────────────────────────────
h("3. Sécurité technique", 1)
bullet("Hébergement dans l'Union européenne : base de données Supabase (PostgreSQL) "
       "en région eu-west-1 (Irlande) ; application sur Vercel.", bold_prefix="Hébergement UE — ")
bullet("Chiffrement en transit (HTTPS/TLS partout) et au repos (chiffrement "
       "Supabase au niveau base et stockage).", bold_prefix="Chiffrement — ")
bullet("Authentification gérée par Clerk (sessions, JWT) ; les routes "
       "non publiques sont protégées par un middleware (proxy.ts) — un visiteur "
       "non authentifié ne peut accéder à aucune donnée prospect.", bold_prefix="Authentification — ")
bullet("Row Level Security (RLS) activée sur les tables sensibles ; les tables "
       "d'audit et internes n'ont AUCUNE policy (accès réservé au rôle technique "
       "service_role côté serveur). Ni prospect, ni professionnel, ni visiteur "
       "ne peut lire les tables d'audit.", bold_prefix="Cloisonnement (RLS) — ")
bullet("Limitation de débit (rate-limiting) sur les endpoints publics et "
       "sensibles (contact, liste d'attente, décisions, vérification téléphone…) "
       "pour prévenir l'abus et l'énumération.", bold_prefix="Rate-limiting — ")
bullet("Paiements et KYC délégués à Stripe (Connect) : les IBAN sortants et la "
       "vérification d'identité sont gérés par Stripe (environnement certifié "
       "PCI-DSS), pas stockés en clair par BUUPP.", bold_prefix="Paiements — ")
proof("proxy.ts ; lib/rate-limit/check.ts ; migrations RLS supabase/migrations/* ; "
      "app/api/prospect/payout/onboarding/route.ts ; app/rgpd/page.tsx")

# ── 4. Anti-fraude multicouche ───────────────────────────────────────
h("4. Anti-fraude multicouche", 1)
para("Mesure affichée publiquement : « Contraintes d'unicité IBAN / téléphone / "
     "rôle, honeypots sur formulaires publics, journal d'audit verrouillé des "
     "révélations. » Voici comment la démontrer :")
bullet("Index UNIQUE sur l'IBAN — un même IBAN ne peut être enregistré que par un "
       "seul compte (empêche le multi-comptes pour cumuler des primes).",
       bold_prefix="Unicité IBAN — ")
proof("supabase/migrations/20260507150000_prospect_rib_unique_iban.sql")
bullet("Index UNIQUE partiel sur le téléphone (numéros non-NULL) — un numéro = un "
       "seul prospect ; toute réutilisation est rejetée (409 phone_already_used). "
       "Le numéro est par ailleurs vérifié par OTP.",
       bold_prefix="Unicité téléphone — ")
proof("supabase/migrations/20260507160000_prospect_identity_unique_phone.sql ; "
      "20260505070000_phone_verification.sql")
bullet("Trigger d'exclusivité de rôle — un même utilisateur ne peut pas être à la "
       "fois prospect et professionnel (violation 23505 → 409).",
       bold_prefix="Exclusivité de rôle — ")
proof("supabase/migrations/20260508140000_role_exclusivity.sql")
bullet("Honeypots — champ caché « website » sur les formulaires publics "
       "(contact, contact DPO, liste d'attente). Rempli = bot → soumission "
       "rejetée (événement waitlist.honeypot_blocked tracé).",
       bold_prefix="Honeypots — ")
proof("app/_components/HomeContactSection.tsx ; app/contact-dpo/_components/ContactDpoForm.tsx ; "
      "app/api/waitlist/route.ts")
bullet("Watermarking cryptographique — l'e-mail révélé au pro est un alias unique "
       "par relation (prospect+rXXX@buupp.com, routé via Cloudflare) : toute fuite "
       "remonte au professionnel émetteur.",
       bold_prefix="Watermarking — ")
proof("lib/aliases/relation-email.ts")
bullet("Journal d'audit verrouillé des révélations — voir section 5.",
       bold_prefix="Journal verrouillé — ")

# ── 5. Journalisation & traçabilité ──────────────────────────────────
h("5. Journalisation et traçabilité des accès", 1)
para("Toute révélation d'une donnée personnelle (e-mail, téléphone, nom complet, "
     "ou ouverture de la fiche détaillée) par un professionnel est journalisée "
     "dans la table pro_contact_reveals : QUI (pro), QUOI (champ), QUEL prospect "
     "(relation), QUAND (horodatage).")
bullet("Écriture FAIL-CLOSED (modifié le 14/06/2026) : la donnée n'est livrée que "
       "si la révélation a pu être journalisée. En cas d'échec d'écriture du "
       "journal, l'API renvoie 500 et n'expose RIEN. Aucune révélation ne peut "
       "donc avoir lieu sans trace.", bold_prefix="Garantie d'exhaustivité — ")
proof("app/api/pro/contacts/[relationId]/reveal/route.ts ; .../details/route.ts ; "
      "app/api/pro/contacts/group-reveal/route.ts")
bullet("Verrou append-only (appliqué en prod le 14/06/2026) : un trigger "
       "PostgreSQL rejette tout UPDATE sur le journal, y compris pour le rôle "
       "technique de l'application → les entrées sont immuables, impossibles à "
       "falsifier a posteriori.", bold_prefix="Inviolabilité — ")
proof("supabase/migrations/20260718120000_pro_contact_reveals_append_only.sql "
      "(trigger pro_contact_reveals_lock_update, base de prod 'buupp')")
bullet("Conservation limitée : 24 mois (durée de politique à valider avec le DPO), "
       "puis purge quotidienne automatique des entrées au-delà. La suppression "
       "par cascade reste possible pour l'exercice du droit à l'effacement.",
       bold_prefix="Conservation — ")
proof("lib/pro/reveals-retention.ts ; cron app/api/admin/digest/route.ts")
para("Autres journaux : clics sur les icônes de contact (pro_contact_clicks, avec "
     "alerte automatique au pro en cas d'abus ≥ 3 clics/24h sur un même prospect) ; "
     "journal d'événements administrateur (admin_events) pour la supervision.")
proof("supabase/migrations/20260604120000_pro_contact_clicks.sql ; "
      "lib/pro/contact-click-alert.ts ; lib/admin/events/record.ts")

# ── 6. Droits des personnes ──────────────────────────────────────────
h("6. Droits des personnes concernées", 1)
bullet("Droit d'accès (art. 15) : le prospect peut savoir quels professionnels "
       "ont accédé à ses données et quand (journal pro_contact_reveals).")
bullet("Droit de rectification (art. 16) et d'effacement (art. 17) : exerçables "
       "depuis l'espace prospect ou via le DPO ; l'effacement supprime les "
       "données et, par cascade, les journaux associés.")
bullet("Transparence (art. 12-14) : politique RGPD publique, page DPO dédiée, "
       "information cookies.")
proof("Pages publiques /rgpd, /contact-dpo, /cookies, /minimisation ; "
      "espace prospect (Prospect.jsx)")

# ── 7. Conservation des données ──────────────────────────────────────
h("7. Limitation de la conservation", 1)
para("Le journal d'audit des révélations est conservé 24 mois puis purgé "
     "automatiquement (cf. section 5). Les données d'un prospect sont supprimées "
     "à l'exercice de son droit à l'effacement, entraînant par cascade la "
     "suppression des relations et journaux liés.")
para("Point de vigilance restant : formaliser au registre des traitements la "
     "durée de conservation retenue pour chaque catégorie de données (compte, "
     "relations, journaux) après validation par le DPO.", italic=True)

# ── 8. Sous-traitants / hébergeurs ───────────────────────────────────
h("8. Sous-traitants et hébergeurs", 1)
bullet("Supabase — base de données PostgreSQL, région UE (Irlande).")
bullet("Vercel — hébergement de l'application.")
bullet("Clerk — authentification / gestion des identités.")
bullet("Stripe — paiements, KYC et virements (Connect), certifié PCI-DSS.")
bullet("Cloudflare — routage des alias e-mail watermarqués.")
bullet("Brevo — envoi des e-mails et SMS transactionnels.")
para("Des accords de sous-traitance (DPA) doivent être tenus à jour avec chacun "
     "de ces prestataires et annexés au registre des traitements.", italic=True)

# ── 9. Tableau de synthèse des preuves ───────────────────────────────
doc.add_page_break()
h("9. Tableau de synthèse — mesure → preuve", 1)
evidence_table([
    ["Pseudonymisation à la lecture", "Masquage/généralisation/suppression appliqués côté serveur",
     "lib/pro/pseudonymize.ts"],
    ["E-mail watermarqué (alias)", "Alias unique par relation, traçable",
     "lib/aliases/relation-email.ts"],
    ["Unicité IBAN", "Index UNIQUE", "20260507150000_prospect_rib_unique_iban.sql"],
    ["Unicité téléphone", "Index UNIQUE partiel + OTP",
     "20260507160000_prospect_identity_unique_phone.sql"],
    ["Exclusivité de rôle", "Trigger 23505", "20260508140000_role_exclusivity.sql"],
    ["Honeypots", "Champ caché 'website' sur formulaires publics",
     "ContactDpoForm.tsx ; HomeContactSection.tsx ; waitlist"],
    ["Journal des révélations", "qui/quoi/quand", "table pro_contact_reveals"],
    ["Audit fail-closed", "Pas de journal → pas de révélation (500)",
     ".../reveal|details|group-reveal/route.ts"],
    ["Journal append-only", "Trigger anti-UPDATE (immuable)",
     "20260718120000_pro_contact_reveals_append_only.sql"],
    ["Rétention 24 mois", "Purge quotidienne", "lib/pro/reveals-retention.ts"],
    ["RLS", "Tables d'audit sans policy (service_role only)", "migrations RLS"],
    ["Rate-limiting", "Endpoints publics/sensibles", "lib/rate-limit/check.ts"],
    ["Authentification", "Clerk + middleware", "proxy.ts"],
    ["Hébergement UE", "Supabase eu-west-1 + Vercel", "config projet"],
])

# ── 10. Changelog ────────────────────────────────────────────────────
h("10. Modifications du 14/06/2026", 1)
bullet("Journalisation des révélations rendue FAIL-CLOSED sur les 3 endpoints "
       "(reveal, details, group-reveal) : aucune révélation sans trace.")
bullet("Verrou append-only ajouté et APPLIQUÉ EN PRODUCTION sur le journal "
       "pro_contact_reveals (trigger anti-UPDATE).")
bullet("Politique de conservation : rétention 24 mois + purge quotidienne "
       "automatique du journal.")

doc.add_paragraph()
foot = doc.add_paragraph()
r = foot.add_run("Document généré automatiquement à partir du code source de "
                 "production. Les emplacements indiqués (« Preuve ») sont "
                 "directement consultables dans le dépôt et la base de données.")
r.italic = True
r.font.size = Pt(8.5)
r.font.color.rgb = GREY

out = "docs/rgpd/BUUPP-conformite-securite-RGPD.docx"
doc.save(out)
print("OK ->", out)
