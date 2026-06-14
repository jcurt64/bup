#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère TOUS les documents de conformité RGPD de BUUPP, chacun en .docx ET .pdf,
à partir d'un MODÈLE DE CONTENU COMMUN (une seule source → deux formats, pas de
divergence). Ancré sur le système réel (politique /rgpd publiée + base + code),
aligné sur les référentiels CNIL.

Documents :
  1. BUUPP-conformite-securite-RGPD        (synthèse sécurité)
  2. Registre-des-traitements-BUUPP        (RoPA — art. 30)
  3. Referentiel-durees-conservation-BUUPP (cycle de vie CNIL)
  4. AIPD-BUUPP-mise-en-relation-scoring   (AIPD — art. 35)
  5. AIPD-BUUPP-paiement-remuneration      (AIPD — art. 35, volet paiement)

Usage : python3 scripts/gen-rgpd-all.py
"""
import os

# ── python-docx ──────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── reportlab ────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, PageBreak, ListFlowable, ListItem)
from reportlab.lib.enums import TA_CENTER

OUT = "docs/rgpd"
DATE = "14/06/2026"
EDITEUR = ("Majelink (éditeur de la plateforme BUUPP) — RCS Pau 892 514 167 — "
           "12 Impasse des Étriers, 64140 Lons")
DPO = "Chargé à la protection des données : dp.buupp@buupp.com"

INDIGO = (0x4F, 0x46, 0xE5)
INK = (0x0F, 0x17, 0x2A)
GREY = (0x47, 0x55, 0x65)
RED = (0xB4, 0x53, 0x09)

DISC = ("Document de travail ancré sur le système en production (politique de "
        "confidentialité publiée, base de données, code source). Les arbitrages "
        "juridiques (bases légales, durées, désignation du DPO) doivent être "
        "validés par le DPO / le conseil juridique. Les champs marqués "
        "« [À COMPLÉTER] » requièrent une décision du responsable de traitement. "
        "Ce document ne constitue pas un avis juridique.")

# Blocs : ("cover", title, sub) | ("h", lvl, txt) | ("p", txt, opts) |
#         ("b", prefix|None, txt) | ("table", headers, rows) | ("pb",) | ("disc",)


# ════════════════════════════════════════════════════════════════════
# RENDU DOCX
# ════════════════════════════════════════════════════════════════════
def rgb(t):
    return RGBColor(*t)

def render_docx(blocks, path):
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10); n.font.color.rgb = rgb(INK)
    for blk in blocks:
        k = blk[0]
        if k == "cover":
            _, title, sub = blk
            t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = t.add_run("BUUPP"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = rgb(INDIGO)
            s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = s.add_run(title); r.font.size = Pt(16)
            s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = s2.add_run(sub); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = rgb(GREY)
            doc.add_paragraph()
            m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
            m.add_run("Responsable de traitement : ").bold = True
            m.add_run(EDITEUR + "\n"); m.add_run(DPO + "\n"); m.add_run("Version : " + DATE)
        elif k == "disc":
            p = doc.add_paragraph(); r = p.add_run(DISC)
            r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = rgb(RED)
        elif k == "h":
            _, lvl, txt = blk
            p = doc.add_heading(txt, level=lvl)
            for r in p.runs:
                r.font.color.rgb = rgb(INDIGO) if lvl <= 1 else rgb(INK)
        elif k == "p":
            txt = blk[1]; opts = blk[2] if len(blk) > 2 else {}
            p = doc.add_paragraph(); r = p.add_run(txt)
            r.italic = opts.get("italic", False); r.bold = opts.get("bold", False)
            if opts.get("color"): r.font.color.rgb = rgb(opts["color"])
            if opts.get("size"): r.font.size = Pt(opts["size"])
        elif k == "b":
            _, prefix, txt = blk
            p = doc.add_paragraph(style="List Bullet")
            if prefix: p.add_run(prefix).bold = True
            p.add_run(txt)
        elif k == "table":
            _, headers, rows = blk
            t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, label in zip(t.rows[0].cells, headers):
                rr = c.paragraphs[0].add_run(label); rr.bold = True; rr.font.size = Pt(8)
            for row in rows:
                cells = t.add_row().cells
                for c, val in zip(cells, row):
                    rr = c.paragraphs[0].add_run(val); rr.font.size = Pt(8)
        elif k == "pb":
            doc.add_page_break()
    doc.save(path)


# ════════════════════════════════════════════════════════════════════
# RENDU PDF
# ════════════════════════════════════════════════════════════════════
def C(t):
    return colors.Color(t[0] / 255, t[1] / 255, t[2] / 255)

def render_pdf(blocks, path):
    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=9.5, leading=13,
                          textColor=C(INK), spaceAfter=5)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=15, leading=18,
                        textColor=C(INDIGO), spaceBefore=12, spaceAfter=6)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontSize=12, leading=15,
                        textColor=C(INK), spaceBefore=8, spaceAfter=4)
    disc = ParagraphStyle("disc", parent=body, fontSize=8, textColor=C(RED), spaceAfter=8)
    ctitle = ParagraphStyle("ct", parent=body, alignment=TA_CENTER, fontSize=28,
                            textColor=C(INDIGO), spaceAfter=8, leading=32)
    csub = ParagraphStyle("cs", parent=body, alignment=TA_CENTER, fontSize=15,
                          textColor=C(INK), spaceAfter=6, leading=18)
    csub2 = ParagraphStyle("cs2", parent=body, alignment=TA_CENTER, fontSize=10.5,
                           textColor=C(GREY), spaceAfter=14)
    cmeta = ParagraphStyle("cm", parent=body, alignment=TA_CENTER, fontSize=9.5)

    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []
    for blk in blocks:
        k = blk[0]
        if k == "cover":
            _, title, sub = blk
            story += [Spacer(1, 30 * mm), Paragraph("BUUPP", ctitle),
                      Paragraph(esc(title), csub), Paragraph(esc(sub), csub2),
                      Spacer(1, 8 * mm),
                      Paragraph("<b>Responsable de traitement :</b> " + esc(EDITEUR), cmeta),
                      Paragraph(esc(DPO), cmeta), Paragraph("Version : " + DATE, cmeta)]
        elif k == "disc":
            story.append(Paragraph(esc(DISC), disc))
        elif k == "h":
            story.append(Paragraph(esc(blk[2]), h1 if blk[1] <= 1 else h2))
        elif k == "p":
            txt = blk[1]; opts = blk[2] if len(blk) > 2 else {}
            st = ParagraphStyle("x", parent=body)
            if opts.get("italic"): st.fontName = "Helvetica-Oblique"
            if opts.get("bold"): st.fontName = "Helvetica-Bold"
            if opts.get("color"): st.textColor = C(opts["color"])
            if opts.get("size"): st.fontSize = opts["size"]
            story.append(Paragraph(esc(txt), st))
        elif k == "b":
            _, prefix, txt = blk
            content = (("<b>" + esc(prefix) + "</b>") if prefix else "") + esc(txt)
            story.append(ListFlowable([ListItem(Paragraph(content, body), value="•")],
                                      bulletType="bullet", start="•", leftIndent=12))
        elif k == "table":
            _, headers, rows = blk
            cell = ParagraphStyle("cell", parent=body, fontSize=7.5, leading=9.5, spaceAfter=0)
            hcell = ParagraphStyle("hcell", parent=cell, textColor=colors.white)
            data = [[Paragraph("<b>" + esc(h) + "</b>", hcell) for h in headers]]
            for row in rows:
                data.append([Paragraph(esc(c), cell) for c in row])
            ncols = len(headers)
            avail = 170 * mm
            tbl = Table(data, colWidths=[avail / ncols] * ncols, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), C(INDIGO)),
                ("GRID", (0, 0), (-1, -1), 0.4, C(GREY)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, C((0xF7, 0xF4, 0xEC))]),
            ]))
            story += [Spacer(1, 3), tbl, Spacer(1, 6)]
        elif k == "pb":
            story.append(PageBreak())

    SimpleDocTemplate(path, pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm,
                      topMargin=18 * mm, bottomMargin=18 * mm).build(story)


def emit(name, blocks):
    os.makedirs(OUT, exist_ok=True)
    render_docx(blocks, f"{OUT}/{name}.docx")
    render_pdf(blocks, f"{OUT}/{name}.pdf")
    print("OK ->", name, "(.docx + .pdf)")


# ════════════════════════════════════════════════════════════════════
# DONNÉES PARTAGÉES
# ════════════════════════════════════════════════════════════════════
SOUS_TRAITANTS = [
    ["Supabase Inc.", "Base de données PostgreSQL + Storage", "UE — Irlande (eu-west-1)", "Art. 28 — DPA"],
    ["Vercel Inc.", "Hébergement de l'application", "États-Unis (servi en UE)", "SCC (clauses types)"],
    ["Clerk Inc.", "Authentification / identités", "États-Unis", "SCC (clauses types)"],
    ["Stripe Payments Europe Ltd", "Paiements, KYC, virements (Connect)", "UE — Irlande", "RT distinct (PCI-DSS)"],
    ["Brevo (Sendinblue SAS)", "E-mails et SMS transactionnels", "UE — France", "Art. 28 — DPA"],
    ["Cloudflare", "Routage des alias e-mail watermarqués", "Réseau mondial (UE)", "SCC (clauses types)"],
    ["DGFiP", "Transmission fiscale annuelle (DAC7)", "France", "Destinataire légal"],
]


# ════════════════════════════════════════════════════════════════════
# DOC 1 — Synthèse sécurité
# ════════════════════════════════════════════════════════════════════
def doc_conformite():
    b = [("cover", "Mesures de sécurité et de protection des données personnelles",
          "Dossier de conformité RGPD — à présenter en cas de contrôle (CNIL ou autre autorité)"),
         ("disc",), ("pb",)]
    b += [("h", 1, "0. Objet du document"),
          ("p", "BUUPP est une plateforme française de mise en relation rémunérée à double "
                "consentement entre particuliers (« prospects ») et professionnels. Les "
                "particuliers acceptent, sollicitation par sollicitation, d'être contactés en "
                "échange d'une rémunération ; aucune donnée n'est transmise sans accord explicite."),
          ("p", "Ce document recense les mesures techniques et organisationnelles de sécurité et "
                "de protection des données, et indique où les vérifier dans le système (fichier de "
                "code, migration, configuration)."),
          ("p", "Principe directeur : protection des données dès la conception (privacy by design) "
                "et par défaut. Expertise RGPD portée par le fondateur, juriste spécialisé en "
                "protection des données personnelles.", {"italic": True})]
    b += [("h", 1, "1. Pseudonymisation et minimisation"),
          ("p", "Les données prospect sont stockées en base mais jamais transmises telles quelles : "
                "la transformation est appliquée à la lecture, côté serveur."),
          ("b", None, "nom → masquage ; e-mail → alias watermarqué ; date de naissance → tranche "
                      "d'âge de 5 ans ; code postal → département ; adresse/poste/revenus → "
                      "suppression ; véhicule/animaux/foyer → catégorisation."),
          ("b", None, "Seul BUUPP peut relier un profil pseudonymisé à la personne (réversible par buupp seul)."),
          ("p", "Preuve : lib/pro/pseudonymize.ts ; app/api/pro/contacts/[relationId]/details/route.ts ; "
                "minimisation par finalité (page /minimisation).", {"size": 8.5, "color": GREY})]
    b += [("h", 1, "2. Consentement (double consentement)"),
          ("p", "Deux consentements explicites : (1) acceptation des CGU à l'inscription, "
                "(2) acceptation explicite de chaque sollicitation. Chaque accord est limité à une "
                "seule sollicitation — pas de réutilisation ni de revente."),
          ("p", "Preuve : lib/cnil/consent.ts ; CGU /cgu, /cgv.", {"size": 8.5, "color": GREY})]
    b += [("h", 1, "3. Sécurité technique"),
          ("b", "Hébergement UE — ", "Supabase (PostgreSQL) en Irlande (eu-west-1) ; application sur Vercel."),
          ("b", "Chiffrement — ", "TLS en transit, chiffrement au repos (Supabase)."),
          ("b", "Authentification — ", "Clerk + middleware (proxy.ts) ; aucune donnée prospect accessible sans authentification."),
          ("b", "Cloisonnement (RLS) — ", "Row Level Security ; tables d'audit sans policy (service_role uniquement)."),
          ("b", "Rate-limiting — ", "endpoints publics/sensibles (lib/rate-limit/check.ts)."),
          ("b", "Paiements — ", "délégués à Stripe Connect (PCI-DSS) ; aucune donnée carte stockée par BUUPP.")]
    b += [("h", 1, "4. Anti-fraude multicouche"),
          ("p", "Mesure publique : « Contraintes d'unicité IBAN / téléphone / rôle, honeypots sur "
                "formulaires publics, journal d'audit verrouillé des révélations. » Démonstration :"),
          ("b", "Unicité IBAN — ", "index UNIQUE (20260507150000_prospect_rib_unique_iban.sql)."),
          ("b", "Unicité téléphone — ", "index UNIQUE partiel + OTP (20260507160000_… ; 20260505070000_phone_verification.sql)."),
          ("b", "Exclusivité de rôle — ", "trigger 23505 (20260508140000_role_exclusivity.sql)."),
          ("b", "Honeypots — ", "champ caché « website » (ContactDpoForm.tsx, HomeContactSection.tsx, /api/waitlist)."),
          ("b", "Watermarking — ", "alias e-mail unique par relation routé via Cloudflare (lib/aliases/relation-email.ts)."),
          ("b", "Journal verrouillé — ", "cf. section 5.")]
    b += [("h", 1, "5. Journalisation et traçabilité des accès"),
          ("p", "Toute révélation (e-mail, téléphone, nom, fiche détaillée) par un pro est "
                "journalisée dans pro_contact_reveals : qui / quoi / quel prospect / quand."),
          ("b", "Exhaustivité (fail-closed) — ", "la donnée n'est livrée que si la révélation est journalisée ; sinon HTTP 500, rien n'est exposé."),
          ("b", "Inviolabilité (append-only) — ", "trigger anti-UPDATE appliqué en prod (pro_contact_reveals_lock_update)."),
          ("b", "Conservation — ", "24 mois puis purge quotidienne (lib/pro/reveals-retention.ts)."),
          ("p", "Autres journaux : clics de contact (pro_contact_clicks + alerte abus) ; événements "
                "administrateur (admin_events).", {"size": 9})]
    b += [("h", 1, "6. Droits des personnes"),
          ("b", None, "Accès (art. 15) : le prospect peut savoir qui a accédé à ses données et quand."),
          ("b", None, "Rectification (16) / effacement (17) : depuis l'espace prospect ou via le DPO ; l'effacement supprime données et journaux liés (cascade)."),
          ("b", None, "Transparence : politique RGPD, page DPO, information cookies.")]
    b += [("h", 1, "7. Sous-traitants et hébergeurs"),
          ("table", ["Sous-traitant", "Rôle", "Localisation", "Encadrement"], SOUS_TRAITANTS)]
    b += [("h", 1, "8. Synthèse — mesure → preuve"),
          ("table", ["Mesure", "Mise en œuvre", "Emplacement (preuve)"], [
              ["Pseudonymisation à la lecture", "Masquage/généralisation/suppression côté serveur", "lib/pro/pseudonymize.ts"],
              ["E-mail watermarqué", "Alias unique par relation", "lib/aliases/relation-email.ts"],
              ["Unicité IBAN/téléphone/rôle", "Index UNIQUE + trigger 23505", "migrations 2026050715/16/08…"],
              ["Honeypots", "Champ caché 'website'", "ContactDpoForm.tsx ; HomeContactSection.tsx"],
              ["Journal des révélations", "qui/quoi/quand", "table pro_contact_reveals"],
              ["Audit fail-closed", "Pas de journal → pas de révélation", ".../reveal|details|group-reveal"],
              ["Journal append-only", "Trigger anti-UPDATE", "20260718120000_…_append_only.sql"],
              ["Rétention 24 mois", "Purge quotidienne", "lib/pro/reveals-retention.ts"],
              ["RLS / Rate-limit / Auth", "service_role only ; limites ; Clerk", "proxy.ts ; lib/rate-limit/check.ts"],
          ])]
    return "BUUPP-conformite-securite-RGPD", b


# ════════════════════════════════════════════════════════════════════
# DOC 2 — Registre des traitements
# ════════════════════════════════════════════════════════════════════
FICHES = [
    ("T1 — Gestion des comptes utilisateurs",
     "Création / gestion du compte, vérification e-mail et téléphone, sécurité.",
     "Exécution du contrat (art. 6.1.b) ; vérification téléphone : intérêt légitime / sécurité.",
     "Prospects, professionnels.",
     "Authentification (id Clerk, e-mail, méthode de connexion) ; Palier 1 identité prospect "
     "(prénom, nom, e-mail, téléphone vérifié SMS, date de naissance, genre, nationalité) ; "
     "profil pro (raison sociale, SIREN/SIRET, forme juridique, RCS, adresse).",
     "Personnels habilités Majelink ; Clerk ; Brevo ; Stripe (KYC pro).",
     "Durée du compte ; suppression des données identifiantes à la clôture.",
     "Téléphone vérifié par OTP ; unicité téléphone/IBAN/e-mail."),
    ("T2 — Données de profil par paliers (2 à 5)",
     "Enrichissement facultatif du profil pour la mise en relation ciblée.",
     "Consentement (art. 6.1.a) — paliers 2 à 5 facultatifs.",
     "Prospects.",
     "Localisation (adresse, CP, ville, région, logement, mobilité) ; style de vie (famille, "
     "véhicule, sport, animaux) ; données professionnelles (poste, revenus, statut, secteur) ; "
     "patrimoine & projets (propriétaire/locataire, épargne, projet immobilier, succession, "
     "création d'entreprise).",
     "Personnels habilités Majelink ; transmis pseudonymisés aux pros (cf. T4).",
     "Durée du compte ; effacement au retrait du consentement.",
     "Données financières (revenus, patrimoine) = hautement personnelles → AIPD."),
    ("T3 — Mise en relation rémunérée (double consentement)",
     "Ciblage par palier, acceptation des sollicitations, calcul de la rémunération, traçabilité.",
     "Exécution du contrat (6.1.b) ; acceptation : consentement (6.1.a).",
     "Prospects, professionnels.",
     "Sollicitations, acceptations/refus, motif du pro, date de décision, code de campagne, "
     "évaluation post-contact.",
     "Prospect, professionnel concerné, personnels Majelink.",
     "Durée du compte / archivage comptable (10 ans).",
     "Aucune donnée sans accord explicite ; séquestre jusqu'à clôture."),
    ("T4 — Pseudonymisation et transmission aux pros",
     "Transmettre un profil exploitable mais non directement identifiant.",
     "Exécution du contrat (6.1.b) ; consentement à la sollicitation (6.1.a).",
     "Prospects.",
     "Profil pseudonymisé : nom masqué, e-mail → alias, naissance → tranche d'âge, CP → "
     "département, adresse/poste/revenus → supprimés.",
     "Professionnel destinataire (RT distinct pour les coordonnées révélées).",
     "Journal des révélations : 24 mois.",
     "Réversible par BUUPP seul ; chaque révélation journalisée (fail-closed + append-only)."),
    ("T5 — Calcul du BUUPP Score et segmentation",
     "Noter la qualité de profil et permettre la segmentation d'audience.",
     "Intérêt légitime (6.1.f). [À COMPLÉTER : balance des intérêts]",
     "Prospects.",
     "Score (0-1000) dérivé des paliers ; facettes (score, région, distance, logement, "
     "statut pro, foyer, véhicule, animaux).",
     "Professionnels (vue agrégée + segments) ; personnels Majelink.",
     "Durée du compte.",
     "Profilage + croisement → AIPD (art. 35)."),
    ("T6 — Gestion financière et rémunération",
     "Recharge des portefeuilles pro, séquestre, débit, paiement et retrait sur IBAN prospect.",
     "Exécution du contrat (6.1.b).",
     "Prospects, professionnels.",
     "Identifiant client Stripe (token), IBAN prospect, historique des transactions wallet, factures.",
     "Stripe (paiements/virements), personnels Majelink.",
     "IBAN : durée du compte ; pièces comptables : 10 ans.",
     "Aucune donnée carte stockée par BUUPP (Stripe, PCI-DSS) → AIPD paiement."),
    ("T7 — Facturation et obligations fiscales (DAC7)",
     "Factures et déclaration DGFiP au-delà des seuils (2 000 € ou 30 transactions/an).",
     "Obligation légale (6.1.c — CGI art. 242 bis, DAC7).",
     "Prospects (au-delà du seuil), professionnels.",
     "Cumul annuel des gains, nombre de transactions, attestations DGFiP, factures.",
     "DGFiP ; personnels Majelink ; comptable.",
     "Fiscal DAC7 : 6 ans ; comptabilité : 10 ans.",
     "—"),
    ("T8 — Communication transactionnelle (e-mail / SMS)",
     "Confirmations, alertes d'encaissement, notifications, OTP téléphone.",
     "Exécution du contrat (6.1.b) ; SMS OTP : sécurité.",
     "Prospects, professionnels.",
     "E-mail, téléphone, contenu transactionnel.",
     "Brevo (UE) ; personnels Majelink.",
     "Durée du compte ; e-mails Pro→Prospect : 12 mois.",
     "Domaine authentifié SPF/DKIM/DMARC."),
    ("T9 — Broadcasts pros → prospects + mesure d'ouverture",
     "Diffusion médiée à un segment ; mesure d'ouverture (pixel).",
     "Exécution du contrat (6.1.b) ; pixel : consentement (6.1.a).",
     "Prospects.",
     "Sujet/corps des messages, identifiant destinataire opaque, statut d'ouverture.",
     "Brevo ; professionnel émetteur ; personnels Majelink.",
     "Tracking pixel : 13 mois ; e-mails : 12 mois.",
     "Pixel inséré uniquement avec consentement explicite."),
    ("T10 — Prévention de la fraude et journalisation",
     "Anti-doublon, exclusivité de rôle, détection d'abus, traçabilité des révélations.",
     "Intérêt légitime (6.1.f) — sécurité et lutte contre la fraude.",
     "Prospects, professionnels.",
     "Contraintes d'unicité (IBAN, téléphone E.164, e-mail, rôle), honeypots, alias "
     "watermarqués, journal des révélations, clics de contact.",
     "Personnels habilités Majelink (conformité).",
     "Journal des révélations : 24 mois ; clics de contact : 24 mois.",
     "Journal verrouillé (append-only) + écriture fail-closed."),
    ("T11 — Liste d'attente (waitlist) et anti-bot",
     "Inscription avant lancement, protection anti-bot.",
     "Consentement (6.1.a) ; anti-bot : intérêt légitime (6.1.f).",
     "Visiteurs / futurs prospects.",
     "E-mail, hash IP (SHA-256 salé, non réversible), user-agent.",
     "Personnels Majelink.",
     "IP hashées : 12 mois ; e-mails de prospection : 3 ans (CNIL).",
     "Honeypot anti-bot ; IP pseudonymisée."),
    ("T12 — Gestion des demandes RGPD / relation DPO",
     "Réception et traitement des demandes d'exercice de droits.",
     "Obligation légale (6.1.c — art. 12 à 22 RGPD).",
     "Prospects, professionnels, anciens utilisateurs, tiers.",
     "Identité du demandeur, objet, échanges, justificatif d'identité.",
     "DPO / personnels habilités Majelink.",
     "Le temps du traitement + délai de preuve. [À COMPLÉTER]",
     "Formulaire DPO public (honeypot + consentement)."),
    ("T13 — Analyse interne et statistiques agrégées",
     "Pilotage du service via des agrégats anonymisés ; back-office d'administration.",
     "Intérêt légitime (6.1.f).",
     "Prospects, professionnels (agrégés).",
     "Agrégats anonymisés ; journal d'événements administrateur.",
     "Personnels habilités Majelink.",
     "Logs techniques/sécurité : 12 mois.",
     "Mesure d'audience via Vercel Analytics — sans cookie."),
]

def doc_registre():
    b = [("cover", "Registre des activités de traitement", "RoPA — article 30 du RGPD"),
         ("disc",), ("pb",)]
    b += [("h", 1, "Informations générales"),
          ("b", "Responsable de traitement : ", EDITEUR),
          ("b", "DPO : ", "dp.buupp@buupp.com"),
          ("b", "Catégories de personnes : ", "Particuliers (« prospects »), professionnels clients, visiteurs (liste d'attente)."),
          ("b", "Transferts hors UE : ", "Données de production hébergées dans l'UE (Supabase, Irlande). "
                "Certains sous-traitants (Vercel, Clerk, Cloudflare) sont aux États-Unis → clauses contractuelles types (SCC).")]
    b += [("h", 1, "Sous-traitants et destinataires (art. 28 / 30.1.d)"),
          ("table", ["Sous-traitant", "Rôle", "Localisation", "Encadrement"], SOUS_TRAITANTS),
          ("p", "Les professionnels destinataires d'une sollicitation acceptée deviennent "
                "responsables de traitement distincts pour les seules coordonnées révélées. "
                "Tenir à jour les DPA (art. 28). [À COMPLÉTER : dates de signature des DPA]", {"italic": True, "size": 9})]
    b += [("h", 1, "Fiches de traitement"),
          ("p", "Mesures de sécurité communes : hébergement UE, chiffrement (TLS + au repos), "
                "authentification Clerk + middleware, RLS, rate-limiting, pseudonymisation à la "
                "lecture, journalisation des accès (fail-closed + append-only), anti-fraude.", {"size": 9})]
    for (titre, fin, base, pers, donnees, dest, duree, part) in FICHES:
        b.append(("h", 2, titre))
        b.append(("b", "Finalité : ", fin))
        b.append(("b", "Base légale : ", base))
        b.append(("b", "Personnes concernées : ", pers))
        b.append(("b", "Catégories de données : ", donnees))
        b.append(("b", "Destinataires / sous-traitants : ", dest))
        b.append(("b", "Durée de conservation : ", duree))
        if part != "—":
            b.append(("b", "Particularités / mesures : ", part))
    return "Registre-des-traitements-BUUPP", b


# ════════════════════════════════════════════════════════════════════
# DOC 3 — Référentiel des durées
# ════════════════════════════════════════════════════════════════════
RETENTION = [
    ["Compte & identité (prospect/pro)", "Durée du compte", "Suppression immédiate des données identifiantes à la clôture", "Nécessité au contrat (CNIL)"],
    ["Paliers 2-5 (profil facultatif)", "Durée du compte tant que le consentement est maintenu", "Effacement à la demande / au retrait du consentement", "Consentement (6.1.a) ; minimisation"],
    ["Transactions / pièces comptables / factures", "Durée de la relation", "Archivage 10 ans", "Art. L.123-22 Code de commerce ; LPF L.102 B"],
    ["Données fiscales (DAC7)", "Exercice en cours", "6 ans après la fin de l'exercice", "CGI art. 242 bis ; directive UE DAC7"],
    ["IBAN prospect (rémunération)", "Durée du compte", "Supprimé à la clôture ; preuve via pièce comptable", "CNIL : données bancaires = temps de la prestation"],
    ["Identifiant client Stripe (token)", "Durée du compte", "Supprimé à la clôture", "Donnée carte gérée par Stripe (RT distinct, PCI-DSS)"],
    ["Journal des révélations (pro_contact_reveals)", "Accès en base active", "24 mois puis purge quotidienne", "Accountability (art. 5.2) — durée validée"],
    ["Actions de contact / click-to-call (audit)", "—", "24 mois après l'événement", "Preuve en cas de litige / signalement"],
    ["Logs techniques et de sécurité", "—", "12 mois maximum", "CNIL ~6 mois ; LCEN (données de connexion) jusqu'à 1 an"],
    ["E-mails Pro → Prospect (contenu)", "—", "12 mois", "Audit anti-spam (intérêt légitime)"],
    ["Tracking pixel des broadcasts", "Consentement", "13 mois après envoi", "Aligné durée traceur CNIL"],
    ["IP hashées (waitlist)", "—", "12 mois", "Anti-bot ; IP pseudonymisée (hash salé)"],
    ["E-mails de prospection (non client)", "Jusqu'au lancement / à la réponse", "3 ans à compter du dernier contact", "Recommandation CNIL — prospection"],
    ["Cookies / traceurs non essentiels", "Dépôt soumis à consentement", "Traceur ≤ 13 mois ; infos ≤ 25 mois", "Délibérations CNIL 2020-091/092"],
    ["Mesure d'audience (Vercel Analytics)", "—", "Sans cookie ; agrégats", "Exemptée de consentement si conforme CNIL"],
    ["Demandes d'exercice de droits (DPO)", "Le temps du traitement", "Durée de preuve. [À COMPLÉTER]", "Obligation légale (art. 12 RGPD)"],
    ["Journal d'événements admin (admin_events)", "—", "12 mois (logs)", "Sécurité / supervision (intérêt légitime)"],
]

def doc_referentiel():
    b = [("cover", "Référentiel des durées de conservation",
          "Cycle de vie des données — base active / archivage / suppression"),
         ("disc",), ("pb",)]
    b += [("h", 1, "Méthode"),
          ("p", "Conformément au principe de limitation de la conservation (art. 5.1.e RGPD) et à "
                "la méthode CNIL du cycle de vie des données, chaque catégorie est conservée :"),
          ("b", None, "en BASE ACTIVE le temps strictement nécessaire à la finalité ;"),
          ("b", None, "puis, le cas échéant, en ARCHIVAGE INTERMÉDIAIRE (accès restreint) pour une obligation légale ou un besoin de preuve ;"),
          ("b", None, "puis SUPPRESSION définitive ou anonymisation."),
          ("p", "Les durées reprennent la politique de confidentialité publiée et les "
                "recommandations / obligations applicables (CNIL, Code de commerce, CGI, LCEN).", {"italic": True, "size": 9})]
    b += [("h", 1, "Tableau des durées"),
          ("table", ["Catégorie de données", "Base active", "Archivage / suppression", "Justification & référence"], RETENTION)]
    b += [("h", 1, "Mises en œuvre techniques de la purge"),
          ("b", None, "Journal des révélations : purge quotidienne (lib/pro/reveals-retention.ts, cron /api/admin/digest), 24 mois."),
          ("b", None, "Bascule de durées CNIL : cron + configuration (lib/cnil/bascule.ts)."),
          ("b", None, "Effacement à la clôture : suppression des données identifiantes, archivage comptable des transactions (10 ans)."),
          ("p", "À finaliser : automatiser la purge des e-mails de prospection (3 ans), des logs "
                "(12 mois) et du tracking pixel (13 mois) si non planifiée ; arbitrer la durée des "
                "demandes DPO. [À COMPLÉTER]", {"italic": True, "size": 9, "color": RED})]
    return "Referentiel-durees-conservation-BUUPP", b


# ════════════════════════════════════════════════════════════════════
# DOC 4 — AIPD mise en relation + scoring
# ════════════════════════════════════════════════════════════════════
def doc_aipd_scoring():
    b = [("cover", "Analyse d'impact relative à la protection des données (AIPD)",
          "Traitement : mise en relation rémunérée avec scoring et segmentation — art. 35 RGPD"),
         ("disc",), ("pb",)]
    b += [("h", 1, "1. Faut-il une AIPD ? (critères CNIL / WP248)"),
          ("p", "Une AIPD est requise en cas de risque élevé probable. Le traitement central de "
                "BUUPP remplit plusieurs des 9 critères du CEPD (WP248) — deux suffisent :"),
          ("b", "✔ ", "Évaluation / scoring : calcul du BUUPP Score."),
          ("b", "✔ ", "Croisement de données : segmentation multi-facettes (revenus, patrimoine, localisation, foyer…)."),
          ("b", "✔ ", "Données hautement personnelles : revenus, patrimoine, projets, IBAN."),
          ("b", "✔ ", "Usage innovant : rémunération de la donnée à double consentement."),
          ("b", "✔ ", "Suivi systématique : journalisation des accès et des sollicitations."),
          ("p", "Conclusion : AIPD requise. Le volet paiement fait l'objet d'une AIPD dédiée.")]
    b += [("h", 1, "2. Description du traitement"),
          ("b", "Finalités : ", "mettre en relation des prospects consentants et des professionnels, en valorisant/segmentant les profils, avec rémunération du prospect."),
          ("b", "Données : ", "identité (palier 1), profil enrichi (paliers 2-5 dont revenus/patrimoine), score dérivé, données de relation."),
          ("b", "Acteurs : ", "prospects (personnes concernées), professionnels, personnels Majelink, sous-traitants (cf. registre)."),
          ("b", "Cycle de vie : ", "collecte (consentement par palier) → scoring/segmentation → sollicitation (double consentement) → pseudonymisation → révélation journalisée → rémunération → archivage/suppression."),
          ("b", "Mesures existantes : ", "hébergement UE, chiffrement, RLS, pseudonymisation à la lecture, alias watermarqués, journal fail-closed + append-only, anti-fraude.")]
    b += [("h", 1, "3. Nécessité et proportionnalité"),
          ("b", "Base légale : ", "contrat (mise en relation) ; consentement (paliers, sollicitations) ; obligation légale (fiscal) ; intérêt légitime (fraude). Balance d'intérêts du scoring à formaliser. [À COMPLÉTER]"),
          ("b", "Minimisation : ", "seuls les paliers nécessaires à l'objectif du pro sont autorisés/facturés ; paliers 2-5 facultatifs et révocables."),
          ("b", "Qualité : ", "données saisies par le prospect ; téléphone vérifié OTP ; unicité anti-doublon."),
          ("b", "Durées : ", "conformes au référentiel de conservation."),
          ("b", "Information & droits : ", "politique de confidentialité, page DPO, page minimisation ; double consentement ; droits d'accès, rectification, effacement, opposition, portabilité.")]
    b += [("h", 1, "4. Appréciation des risques pour les personnes"),
          ("p", "Trois événements redoutés (méthode CNIL) : accès illégitime, modification non "
                "désirée, disparition. Cotation gravité × vraisemblance à valider en atelier DPO.", {"italic": True, "size": 9}),
          ("table", ["Risque", "Impact pour la personne", "Mesures en place", "Risque résiduel*"], [
              ["Accès illégitime (ré-identification, fuite IBAN/revenus)", "Atteinte vie privée, démarchage abusif, risque financier",
               "Pseudonymisation à la lecture ; alias watermarqués ; RLS ; chiffrement ; révélation gated + journalisée ; aucune carte stockée", "Limité [à valider]"],
              ["Détournement de finalité par un pro (extraction, revente)", "Sollicitations hors cadre, perte de contrôle",
               "Données non exportables ; consentement à usage unique ; watermark traçable ; journal inviolable ; alerte abus", "Limité [à valider]"],
              ["Falsification d'une trace d'accès", "Impossibilité de prouver l'accès",
               "Journal append-only (trigger anti-UPDATE) + écriture fail-closed", "Négligeable [à valider]"],
              ["Profilage abusif (scoring) / décision défavorable", "Exclusion d'opportunités, traitement inéquitable",
               "Score non décisionnel à effet juridique ; transparence ; pas de catégorie sensible ; opposition/rectification", "Limité [à valider]"],
              ["Disparition / indisponibilité", "Perte d'historique, de rémunération due",
               "Hébergement managé (sauvegardes) ; séquestre des fonds", "Limité [à valider]"],
          ]),
          ("p", "* Cotation finale à arrêter avec le DPO (outil PIA CNIL). [À COMPLÉTER]", {"size": 8.5, "italic": True, "color": RED})]
    b += [("h", 1, "5. Plan d'action et validation"),
          ("b", None, "Formaliser la balance d'intérêts du scoring (intérêt légitime)."),
          ("b", None, "Arbitrer et automatiser les purges manquantes (prospection 3 ans, logs 12 mois)."),
          ("b", None, "Tenir à jour les DPA des sous-traitants."),
          ("b", None, "Valider la cotation des risques (outil PIA CNIL) avec le DPO."),
          ("table", ["Rôle", "Nom", "Avis / décision", "Date"], [
              ["Responsable de traitement", "[À COMPLÉTER]", "", ""],
              ["DPO", "[À COMPLÉTER]", "", ""],
          ])]
    return "AIPD-BUUPP-mise-en-relation-scoring", b


# ════════════════════════════════════════════════════════════════════
# DOC 5 — AIPD paiement / rémunération (NOUVEAU)
# ════════════════════════════════════════════════════════════════════
def doc_aipd_paiement():
    b = [("cover", "Analyse d'impact relative à la protection des données (AIPD)",
          "Traitement : gestion financière et rémunération (wallet, séquestre, IBAN, retraits) — art. 35 RGPD"),
         ("disc",), ("pb",)]
    b += [("h", 1, "1. Faut-il une AIPD ? (critères CNIL / WP248)"),
          ("p", "Le traitement des données financières (coordonnées bancaires, flux de rémunération) "
                "présente un risque élevé potentiel. Critères WP248 remplis :"),
          ("b", "✔ ", "Données hautement personnelles : coordonnées bancaires (IBAN), montants, historique de transactions."),
          ("b", "✔ ", "Collecte à grande échelle de données financières de particuliers rémunérés."),
          ("b", "✔ ", "Usage innovant : versement direct de micro-rémunérations liées à la donnée personnelle."),
          ("p", "Particularité : l'essentiel de la chaîne de paiement (vérification d'identité KYC, "
                "rails de virement, traitement carte) est délégué à Stripe (Connect), responsable de "
                "traitement distinct, certifié PCI-DSS. Le risque résiduel pour BUUPP porte sur le "
                "stockage de l'IBAN de retrait, le grand livre du portefeuille (wallet) et les "
                "factures. Une AIPD ciblée est néanmoins justifiée par la sensibilité des données.")]
    b += [("h", 1, "2. Description du traitement"),
          ("b", "Finalités : ", "recharge des portefeuilles pro, mise sous séquestre, débit à la mise en relation, calcul et versement de la rémunération du prospect, retrait sur IBAN, facturation."),
          ("b", "Données : ", "identifiant client Stripe (token, pas de numéro de carte), IBAN du prospect, grand livre des transactions wallet, montants, factures, données fiscales agrégées (cf. DAC7)."),
          ("b", "Acteurs : ", "prospects (bénéficiaires), professionnels (payeurs), Stripe (sous-traitant/RT distinct), personnels Majelink (comptabilité, support), DGFiP (obligation fiscale)."),
          ("b", "Cycle de vie : ", "saisie de l'IBAN par le prospect → vérification/virement via Stripe → écriture au grand livre → facture → archivage comptable (10 ans) → suppression de l'IBAN à la clôture du compte."),
          ("b", "Mesures existantes : ", "délégation à Stripe (PCI-DSS), aucune donnée carte stockée par BUUPP, chiffrement, RLS, unicité IBAN (anti-doublon), authentification, journalisation administrateur.")]
    b += [("h", 1, "3. Nécessité et proportionnalité"),
          ("b", "Base légale : ", "exécution du contrat (art. 6.1.b) pour le paiement/retrait ; obligation légale (6.1.c) pour la conservation comptable et la déclaration DAC7."),
          ("b", "Minimisation : ", "BUUPP ne stocke que l'IBAN de retrait et un token Stripe — aucune donnée carte complète ; le KYC est porté par Stripe."),
          ("b", "Qualité : ", "IBAN saisi par le prospect ; contrôle d'unicité (un IBAN = un compte) ; vérifications Stripe."),
          ("b", "Durées : ", "IBAN = durée du compte ; pièces comptables = 10 ans (L.123-22) ; fiscal DAC7 = 6 ans."),
          ("b", "Information & droits : ", "politique de confidentialité (section paiement) ; droits d'accès/rectification/effacement, sous réserve des obligations comptables et fiscales (archivage).")]
    b += [("h", 1, "4. Appréciation des risques pour les personnes"),
          ("p", "Événements redoutés (méthode CNIL). Cotation gravité × vraisemblance à valider avec le DPO.", {"italic": True, "size": 9}),
          ("table", ["Risque", "Impact pour la personne", "Mesures en place", "Risque résiduel*"], [
              ["Fuite / accès illégitime à l'IBAN ou aux montants", "Risque financier, atteinte à la vie privée",
               "Chiffrement ; RLS ; accès restreint ; aucune carte stockée ; KYC/rails délégués à Stripe (PCI-DSS)", "Limité [à valider]"],
              ["Détournement (virement frauduleux, usurpation d'IBAN)", "Perte financière directe",
               "Unicité IBAN (anti-doublon) ; vérifications Stripe ; séquestre ; journalisation", "Limité [à valider]"],
              ["Conservation excessive de données bancaires", "Atteinte à la limitation de conservation",
               "IBAN supprimé à la clôture ; seules les pièces comptables sont archivées (10 ans)", "Négligeable [à valider]"],
              ["Erreur de versement / non-paiement", "Préjudice financier pour le prospect",
               "Grand livre wallet ; séquestre jusqu'à clôture ; traçabilité des transactions", "Limité [à valider]"],
              ["Transfert hors UE des données de paiement", "Perte de garanties",
               "Stripe Payments Europe (Irlande, UE) ; encadrement contractuel", "Négligeable [à valider]"],
          ]),
          ("p", "* Cotation finale à arrêter avec le DPO (outil PIA CNIL). [À COMPLÉTER]", {"size": 8.5, "italic": True, "color": RED})]
    b += [("h", 1, "5. Plan d'action et validation"),
          ("b", None, "Confirmer la qualification de Stripe (sous-traitant vs responsable conjoint/distinct) et tenir le DPA à jour."),
          ("b", None, "Vérifier la suppression effective de l'IBAN à la clôture (hors obligations d'archivage)."),
          ("b", None, "Documenter la politique d'archivage comptable (accès restreint, durée 10 ans)."),
          ("b", None, "Valider la cotation des risques (outil PIA CNIL) avec le DPO."),
          ("table", ["Rôle", "Nom", "Avis / décision", "Date"], [
              ["Responsable de traitement", "[À COMPLÉTER]", "", ""],
              ["DPO", "[À COMPLÉTER]", "", ""],
          ])]
    return "AIPD-BUUPP-paiement-remuneration", b


if __name__ == "__main__":
    for builder in (doc_conformite, doc_registre, doc_referentiel,
                    doc_aipd_scoring, doc_aipd_paiement):
        name, blocks = builder()
        emit(name, blocks)
